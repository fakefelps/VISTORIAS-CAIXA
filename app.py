# -*- coding: utf-8 -*-
"""
BERÇAN PROJETOS — Preenchimento Automático de Documentos CAIXA
Versão 4.5 — Abril/2026

Correções v4 (original):
- Assinatura Memorial (Excel) via win32com com qualidade preservada (+50% tamanho)
- Assinatura Declaração (Word) +50% tamanho
- Checkboxes do Memorial via imagens PNG sobrepostas (confiável e editável)
- OCR da ART com detecção expandida + pré-processamento da imagem + regex tolerantes
- Bug do Modo Não Mapeado (float & int) corrigido
- Pasta destino renomeada: DOCUMENTOS DE VISTORIA
- Botão INTERROMPER processamento (threading.Event)

Correções v4.2:
- Assinatura Word: âncora dinâmica no parágrafo ____ (detectado automaticamente)
- Assinatura Word: posOffset -457200 EMU (0,5cm acima da linha de assinatura)
- Excel COM: xl.Visible/DisplayAlerts envolvidos em try/except (fix AttributeError)
- UI: botão LER ART e campo PDF removidos (OCR descartado)
- UF: campo mantido com padrão GO

Correções v4.1:
- CHECKBOX_ANCORA_CELULA corrigida de AR55 → AM70 (posição real no drawing XML)
- CHECKBOX_LARGURA_PT/ALTURA_PT ajustados para cobrir AM70:AP70 corretamente
- asset_checkbox() adicionada: fallback automático de extensão (.png/.jpeg/.png.jpeg)
- ASSINATURA_EXCEL_ANCORA ajustada de AE73 → AE74 (sem offset negativo frágil)
- Assinatura Word: posOffset corrigido para 0/−685800 (alinha à esquerda da coluna)
- SHAPE_ESGOTO_SIM/NAO confirmados via inspeção do drawing1.xml do template real
- UF padrão mantido como GO (OCR descartado — preenchimento manual preferido)
"""

# ============================================================
# PROTEÇÃO ANTI-LOOP PyInstaller — DEVE SER A PRIMEIRA COISA
# ============================================================
import multiprocessing
multiprocessing.freeze_support()

# ============================================================
# IMPORTS
# ============================================================
import os
import tempfile
import sys
import shutil
import zipfile
import threading
import datetime
import re
import traceback
from pathlib import Path
from io import BytesIO
from copy import deepcopy

import tkinter as tk

# ── SCPO: imports Selenium (usados apenas ao clicar "Preencher SCPO") ────────
import winreg as _winreg
import json as _json_scpo
import urllib.request as _urllib_scpo
import zipfile as _zipfile_scpo
from selenium import webdriver as _webdriver_scpo
from selenium.webdriver.common.by import By as _By_scpo
from selenium.webdriver.support.ui import WebDriverWait as _Wait_scpo
from selenium.webdriver.support.ui import Select as _Select_scpo
from selenium.webdriver.support import expected_conditions as _EC_scpo
from selenium.webdriver.chrome.service import Service as _Service_scpo

def buscar_cep(cep, callback_ok, callback_erro):
    """
    Consulta ViaCEP em thread separada e chama callback com o resultado.
    callback_ok(data): dict com logradouro, bairro, localidade, uf
    callback_erro(msg): string com mensagem de erro
    """
    import urllib.request, json, threading, ssl
    _ssl_ctx = ssl.create_default_context()
    _ssl_ctx.check_hostname = False
    _ssl_ctx.verify_mode = ssl.CERT_NONE
    cep_limpo = re.sub(r"[^0-9]", "", cep)
    if len(cep_limpo) != 8:
        callback_erro("CEP deve ter 8 dígitos")
        return
    def _worker():
        try:
            url = f"https://viacep.com.br/ws/{cep_limpo}/json/"
            with urllib.request.urlopen(url, timeout=5, context=_ssl_ctx) as r:
                data = json.loads(r.read().decode())
            if data.get("erro"):
                callback_erro("CEP não encontrado")
            else:
                callback_ok(data)
        except Exception as e:
            callback_erro(f"Erro na consulta: {e}")
    threading.Thread(target=_worker, daemon=True).start()

from tkinter import ttk, filedialog, messagebox

# Word
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.oxml.ns import qn

# Excel + imagens
from PIL import Image, ImageFilter, ImageOps

# COM (Word/Excel nativo)
import win32com.client
import pythoncom

# XML (preservação de namespaces)
from lxml import etree


# ============================================================
# CONSTANTES GLOBAIS
# ============================================================

# ----- Pasta de destino -----
PASTA_DESTINO = "DOCUMENTOS DE VISTORIA"

# ----- Templates Word -----
TEMPLATE_FOSSA = "TEMPLETE PARA FOSSA.docx"
TEMPLATE_ESGOTO = "TEMPLETE PARA ESGOTO.docx"
FOSSA_LINHA_ASS = 36
ESGOTO_LINHA_ASS = 41

# ----- Checkboxes — método IMAGEM (PNG sobreposto via Shapes.AddPicture) -----
# Valores calibrados com o Calibrador do Memorial.
# Para recalibrar: rodar o Calibrador, ajustar até o PDF ficar correto,
# copiar os valores gerados e colar aqui substituindo os blocos abaixo.
#
# 1. Esgoto SIM
CHK1_ANCORA  = "AM70"
CHK1_OFF_X   = 10
CHK1_OFF_Y   = 3
CHK1_LARGURA = 4
CHK1_ALTURA  = 5
# 2. Esgoto NÃO
CHK2_ANCORA  = "AP70"
CHK2_OFF_X   = 11
CHK2_OFF_Y   = 3
CHK2_LARGURA = 4
CHK2_ALTURA  = 5
# 3. Condomínio SIM
CHK3_ANCORA  = "AM65"
CHK3_OFF_X   = 10
CHK3_OFF_Y   = 8
CHK3_LARGURA = 4
CHK3_ALTURA  = 5
# 4. Condomínio Não se aplica
CHK4_ANCORA  = "AS65"
CHK4_OFF_X   = 12
CHK4_OFF_Y   = 8
CHK4_LARGURA = 4
CHK4_ALTURA  = 5

# Loteamentos — fixo NSA, sempre inserido
CHK_LOT_NSA_ANCORA  = "AS64"
CHK_LOT_NSA_OFF_X   = 12
CHK_LOT_NSA_OFF_Y   = 8
CHK_LOT_NSA_LARGURA = 4
CHK_LOT_NSA_ALTURA  = 5

# Estado de condomínios — atualizado pela UI antes do processamento
GEMINADAS_CONDOMINIOS = "nao_se_aplica"

# ----- Assinatura Word (DECLARAÇÃO) -----
ASSINATURA_WORD_LARGURA = Inches(2.7)

# ----- Assinatura Excel (MEMORIAL) — calibrado -----
ASSINATURA_EXCEL_ANCORA      = "AE72"
ASSINATURA_EXCEL_OFFSET_X_PT = 10
ASSINATURA_EXCEL_OFFSET_Y_PT = -5
ASSINATURA_EXCEL_LARGURA_PT  = 170
ASSINATURA_EXCEL_ALTURA_PT   = 55

# ----- Engenheiros cadastrados -----
ENGENHEIROS = {
    "FELIPE GUILHERME BERÇAN": {
        "cpf": "147.849.107-86",
        "crea": "1022722034D-GO",
        "assinatura": "FELIPE.png",
    },
    "CAIO ARAUJO BRAGA": {
        "cpf": "011.309.411-67",
        "crea": "CREA-GO",
        "assinatura": "CAIO.png",
    },
    "JOÃO VITOR CABRAL DE MORAIS": {
        "cpf": "038.144.411-25",
        "crea": "CREA-GO",
        "assinatura": "JOÃO VITOR.jpg",
    },
    "JULIO CESAR GOMES DE MORAIS FILHO": {
        "cpf": "033.865.821-17",
        "crea": "CREA-GO",
        "assinatura": "JULIO CESAR.png",
    },
    "PAULA FLEURY DE MORAIS": {
        "cpf": "033.813.881-18",
        "crea": "CREA-GO",
        "assinatura": "PAULA.png",
    },
    "ISAAC NATAN SANTOS": {
        "cpf": "701.117.261-07",
        "crea": "CREA-GO",
        "assinatura": "ISAAC.png",
    },
}

# ----- Paleta de cores da UI -----
COR_FUNDO = "#1e2a3a"
COR_LOG_FUNDO = "#131c26"
COR_CAMPO = "#2a3f55"
COR_BOTAO = "#2e86de"
COR_BOTAO_STOP = "#e74c3c"
COR_PROGRESSO = "#4cd964"
COR_TEXTO = "#ffffff"
COR_TEXTO_SEC = "#90adc4"
COR_LOG_TEXTO = "#7ec8a0"


# ============================================================
# PERSISTÊNCIA — arquivo de configuração do usuário
# ============================================================
ARQUIVO_CONFIG = Path.home() / ".bercan_config.json"

def _config_carregar() -> dict:
    """Lê o arquivo de config; retorna dict vazio se não existir ou estiver corrompido."""
    try:
        if ARQUIVO_CONFIG.exists():
            import json
            return json.loads(ARQUIVO_CONFIG.read_text(encoding="utf-8"))
    except Exception:
        pass
    return {}

def _config_salvar(dados: dict):
    """Salva o dict de configuração em disco (merge com o que já existe)."""
    try:
        import json
        atual = _config_carregar()
        atual.update(dados)
        ARQUIVO_CONFIG.write_text(
            json.dumps(atual, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )
    except Exception:
        pass


# ============================================================
# HELPERS — CAMINHOS
# ============================================================

def resource_path(rel: str) -> str:
    """
    Resolve caminho de asset, funciona tanto em .py quanto em .exe do PyInstaller.
    O PyInstaller extrai arquivos para _MEIPASS em runtime.
    """
    base = getattr(sys, "_MEIPASS", str(Path(__file__).parent))
    return str(Path(base) / rel)


def asset(nome: str) -> str:
    """Retorna o caminho completo de um arquivo em assets/."""
    return resource_path(os.path.join("assets", nome))


# ============================================================
# HELPERS — DATA
# ============================================================

def formatar_data_hoje() -> str:
    """Retorna data atual no formato DD/MM/AAAA."""
    return datetime.date.today().strftime("%d/%m/%Y")


def formatar_data_extenso() -> str:
    """Retorna data no formato '16 de abril de 2026'."""
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho",
             "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    hoje = datetime.date.today()
    return f"{hoje.day} de {meses[hoje.month - 1]} de {hoje.year}"


# ============================================================
# HELPERS — WORD (python-docx)
# ============================================================

def _preto_run(run):
    """Força a cor de um run para preto (RGB 0,0,0)."""
    run.font.color.rgb = RGBColor(0, 0, 0)


def _preto_paragrafo(para):
    """Força todos os runs do parágrafo para preto."""
    for run in para.runs:
        _preto_run(run)


def _sub_paragrafo(para, placeholder, valor):
    """
    Substitui placeholder em um parágrafo, consolidando runs fragmentados.
    O Word quebra texto em múltiplos runs por formatação (ex: '{6}' pode estar em
    runs separados: '{', '6', '}'). Reconstruímos o texto, substituímos e
    reescrevemos tudo no primeiro run, zerando os demais.
    """
    if placeholder not in para.text:
        return False
    texto_completo = "".join(r.text for r in para.runs)
    if placeholder not in texto_completo:
        return False
    texto_novo = texto_completo.replace(placeholder, str(valor))
    if para.runs:
        para.runs[0].text = texto_novo
        _preto_run(para.runs[0])
        for run in para.runs[1:]:
            run.text = ""
    return True


def _detectar_paragrafo_assinatura(doc):
    """
    Detecta o índice do parágrafo de assinatura no template Word.
    Procura o parágrafo que contém apenas underscores (____).
    Fallback: penúltimo parágrafo antes de 'RT:'.
    """
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        txt = p.text.strip()
        if txt and all(c in ('_', ' ') for c in txt) and len(txt) >= 5:
            return i
    for i, p in enumerate(paras):
        if p.text.strip().startswith("RT:"):
            return max(0, i - 2)
    return max(0, len(paras) - 2)


def _inserir_assinatura_word(doc, img_path, linha_idx=None, log=None):
    """
    Insere a assinatura como imagem FLUTUANTE (behind text) no Word.
    Ancora dinamicamente no parágrafo com ____ (linha de assinatura).
    """
    if not os.path.exists(img_path):
        if log:
            log(f"⚠ Assinatura não encontrada: {img_path}")
        return

    idx = _detectar_paragrafo_assinatura(doc)
    if log:
        log(f"  • Assinatura ancorada no parágrafo [{idx}]: {repr(doc.paragraphs[idx].text[:40])}")

    target = doc.paragraphs[idx]
    run = target.add_run()
    run.add_picture(img_path, width=ASSINATURA_WORD_LARGURA)

    drawing = run._r.find(qn("w:drawing"))
    if drawing is None:
        return
    inline = drawing.find(qn("wp:inline"))
    if inline is None:
        return

    graphic_elems = [c for c in inline if "graphic" in c.tag]
    if not graphic_elems:
        return
    graphic_el = graphic_elems[0]

    extent_el = inline.find(qn("wp:extent"))
    cx = extent_el.get("cx") if extent_el is not None else "1800000"
    cy = extent_el.get("cy") if extent_el is not None else "600000"

    anchor_xml = f'''<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
        xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        distT="0" distB="0" distL="0" distR="0"
        simplePos="0" relativeHeight="251658240"
        behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="column">
            <wp:posOffset>0</wp:posOffset>
        </wp:positionH>
        <wp:positionV relativeFrom="paragraph">
            <wp:posOffset>-457200</wp:posOffset>
        </wp:positionV>
        <wp:extent cx="{cx}" cy="{cy}"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:wrapNone/>
        <wp:docPr id="1" name="Assinatura"/>
        <wp:cNvGraphicFramePr/>
    </wp:anchor>'''

    anchor = etree.fromstring(anchor_xml)
    anchor.append(deepcopy(graphic_el))
    drawing.remove(inline)
    drawing.append(anchor)


def preencher_word(esgoto_sim, saida_path, dados, num_casa, log=None):
    """Preenche o template Word (FOSSA ou ESGOTO) e salva como .docx."""
    tpl_nome = TEMPLATE_ESGOTO if esgoto_sim else TEMPLATE_FOSSA
    tpl_path = asset(tpl_nome)
    linha_ass = ESGOTO_LINHA_ASS if esgoto_sim else FOSSA_LINHA_ASS

    if not os.path.exists(tpl_path):
        raise FileNotFoundError(f"Template Word não encontrado: {tpl_path}")

    doc = Document(tpl_path)

    subs = {
        "{1}": dados.get("art", ""),
        "{2}": dados.get("crea", ""),
        "{5}": dados.get("logradouro", ""),
        "{6}": dados.get("quadra_lote", ""),
        "{7}": dados.get("bairro", ""),
        "{9}": f"CASA {num_casa}",
        "{10}": dados.get("cidade", ""),
        "{11}": dados.get("uf", ""),
        "{ENGENHEIRO SELECIONADO}": dados.get("engenheiro_nome", ""),
        "{dia/mes/ano}": formatar_data_hoje(),
    }

    for para in doc.paragraphs:
        for ph, val in subs.items():
            _sub_paragrafo(para, ph, val)
        _preto_paragrafo(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for ph, val in subs.items():
                        _sub_paragrafo(para, ph, val)
                    _preto_paragrafo(para)

    _inserir_assinatura_word(doc, dados["assinatura_path"], linha_ass, log)

    doc.save(saida_path)
    if log:
        log(f"  ✓ Word gerado: {os.path.basename(saida_path)}")


def _quadrado_preto_temp():
    """Gera PNG temporário de quadrado preto sólido (■)."""
    tmp = tempfile.mktemp(suffix=".png")
    Image.new("RGBA", (20, 20), (0, 0, 0, 255)).save(tmp)
    return tmp


def _inserir_checkbox_img(ws, ancora, off_x, off_y, largura, altura, img_path):
    """Insere quadrado preto em posição calibrada via Shapes.AddPicture."""
    cell = ws.Range(ancora)
    ws.Shapes.AddPicture(
        os.path.abspath(img_path),
        False, True,
        cell.Left + off_x,
        cell.Top  + off_y,
        largura,
        altura,
    )


def _fechar_excel(xl, wb):
    """Fecha wb e xl com segurança — sempre mata EXCEL.EXE."""
    if wb is not None:
        try: wb.Close(SaveChanges=False)
        except Exception: pass
    if xl is not None:
        try: xl.Quit()
        except Exception: pass
        try:
            import subprocess
            subprocess.run(["taskkill", "/F", "/IM", "EXCEL.EXE"],
                           capture_output=True, creationflags=0x08000000)
        except Exception: pass


def _excel_preencher(template_path, xlsx_saida, dados, num_casa, esgoto_sim, log=None):
    """
    Copia o template virgem para xlsx_saida e preenche via win32com.
    Checkboxes inseridos via imagem na mesma sessão COM.
    """
    import pythoncom, win32com.client

    pythoncom.CoInitialize()
    xl = None
    wb = None
    try:
        xl = win32com.client.Dispatch("Excel.Application")
        try: xl.Visible = False
        except Exception: pass
        try: xl.DisplayAlerts = False
        except Exception: pass
        try: xl.ScreenUpdating = False
        except Exception: pass

        wb = xl.Workbooks.Open(os.path.abspath(template_path))

        try:
            ws = wb.Worksheets("ElemConstrutivos")
        except Exception:
            ws = wb.Worksheets(1)
            if log:
                log(f"  ⚠ Aba 'ElemConstrutivos' não encontrada, usando: {ws.Name}")

        mapa = {
            "G40": dados.get("contratante", ""),
            "G43": dados.get("engenheiro_nome", ""),
            "AH43": dados.get("crea", ""),
            "AP43": "GO",
            "AR43": dados.get("cpf", ""),
            "G47": dados.get("logradouro", ""),
            "AJ47": f"{dados.get('quadra_lote', '')}   CASA {num_casa}",
            "G49": dados.get("bairro", ""),
            "V49": dados.get("cep", ""),
            "AA49": dados.get("cidade", ""),
            "AU49": dados.get("uf", ""),
            "H53": dados.get("engenheiro_nome", ""),
            "Y54": dados.get("art", ""),
            "H75": f"GOIÂNIA, {formatar_data_extenso()}",
            "AE77": dados.get("engenheiro_nome", ""),
            "AE78": dados.get("cpf", ""),
            "AE79": dados.get("crea", ""),
        }
        for coord, val in mapa.items():
            try:
                rng = ws.Range(coord)
                rng.Value = val
                rng.Font.Color = 0
            except Exception as e:
                if log:
                    log(f"  ⚠ Célula {coord} falhou: {e}")

        # Inserir assinatura
        ass_path = dados.get("assinatura_path", "")
        if ass_path and os.path.exists(ass_path):
            try:
                cell = ws.Range(ASSINATURA_EXCEL_ANCORA)
                left = cell.Left + ASSINATURA_EXCEL_OFFSET_X_PT
                top  = cell.Top  + ASSINATURA_EXCEL_OFFSET_Y_PT
                ws.Shapes.AddPicture(
                    os.path.abspath(ass_path),
                    False, True,
                    left, top,
                    ASSINATURA_EXCEL_LARGURA_PT,
                    ASSINATURA_EXCEL_ALTURA_PT,
                )
                if log:
                    log("  ✓ Assinatura inserida")
            except Exception as e:
                if log:
                    log(f"  ⚠ Falha ao inserir assinatura: {e}")

        # ── Checkboxes via imagem (método calibrado) ─────────────────────────
        q = _quadrado_preto_temp()
        try:
            # Esgoto SIM ou NÃO
            if esgoto_sim:
                _inserir_checkbox_img(ws, CHK1_ANCORA, CHK1_OFF_X, CHK1_OFF_Y,
                                      CHK1_LARGURA, CHK1_ALTURA, q)
                if log: log(f"  ✓ Checkbox esgoto SIM ({CHK1_ANCORA})")
            else:
                _inserir_checkbox_img(ws, CHK2_ANCORA, CHK2_OFF_X, CHK2_OFF_Y,
                                      CHK2_LARGURA, CHK2_ALTURA, q)
                if log: log(f"  ✓ Checkbox esgoto NÃO ({CHK2_ANCORA})")

            # Condomínios
            if GEMINADAS_CONDOMINIOS == "sim":
                _inserir_checkbox_img(ws, CHK3_ANCORA, CHK3_OFF_X, CHK3_OFF_Y,
                                      CHK3_LARGURA, CHK3_ALTURA, q)
                if log: log(f"  ✓ Checkbox condomínio SIM ({CHK3_ANCORA})")
            elif GEMINADAS_CONDOMINIOS == "nao_se_aplica":
                _inserir_checkbox_img(ws, CHK4_ANCORA, CHK4_OFF_X, CHK4_OFF_Y,
                                      CHK4_LARGURA, CHK4_ALTURA, q)
                if log: log(f"  ✓ Checkbox condomínio NSA ({CHK4_ANCORA})")

            # Loteamentos — sempre NSA (fixo)
            _inserir_checkbox_img(ws, CHK_LOT_NSA_ANCORA, CHK_LOT_NSA_OFF_X,
                                  CHK_LOT_NSA_OFF_Y, CHK_LOT_NSA_LARGURA,
                                  CHK_LOT_NSA_ALTURA, q)
            if log: log(f"  ✓ Checkbox loteamento NSA ({CHK_LOT_NSA_ANCORA})")

        except Exception as e:
            if log: log(f"  ⚠ Falha ao inserir checkboxes: {e}")
        finally:
            try: os.unlink(q)
            except: pass

        # Salvar como .xlsx (51 = xlOpenXMLWorkbook)
        wb.SaveAs(os.path.abspath(xlsx_saida), FileFormat=51)

    finally:
        _fechar_excel(xl, wb)
        pythoncom.CoUninitialize()


def _excel_para_pdf(xlsx_path, pdf_path, log=None):
    """Exporta Excel para PDF em 1 página via win32com."""
    pythoncom.CoInitialize()
    xl = None
    wb = None
    try:
        xl = win32com.client.Dispatch("Excel.Application")
        try: xl.Visible = False
        except Exception: pass
        try: xl.DisplayAlerts = False
        except Exception: pass
        wb = xl.Workbooks.Open(os.path.abspath(xlsx_path))

        try:
            ws = wb.Worksheets("ElemConstrutivos")
        except Exception:
            ws = wb.Worksheets(1)

        ws.PageSetup.Zoom = False
        ws.PageSetup.FitToPagesWide = 1
        ws.PageSetup.FitToPagesTall = 1

        # ExportAsFixedFormat(Type, Filename, Quality, IncludeDocProperties, IgnorePrintAreas)
        ws.ExportAsFixedFormat(
            0,   # xlTypePDF
            os.path.abspath(pdf_path),
            0,   # xlQualityStandard
            True,
            False,
        )
        if log:
            log(f"  ✓ PDF gerado: {os.path.basename(pdf_path)}")

    finally:
        _fechar_excel(xl, wb)
        pythoncom.CoUninitialize()


def _word_para_pdf(docx_path, pdf_path, log=None):
    """Converte Word para PDF via win32com."""
    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        try: word.Visible = False
        except Exception: pass
        try: word.DisplayAlerts = False
        except Exception: pass
        doc = word.Documents.Open(os.path.abspath(docx_path))
        # 17 = wdFormatPDF — SaveAs2 é o método correto em Word 2013+
        # Fallback para SaveAs em versões mais antigas
        try:
            doc.SaveAs2(os.path.abspath(pdf_path), FileFormat=17)
        except AttributeError:
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        if log:
            log(f"  ✓ PDF gerado: {os.path.basename(pdf_path)}")
    finally:
        if doc is not None:
            try:
                doc.Close(SaveChanges=False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


# ============================================================
# OCR DA ART (melhorado v4)
# ============================================================

def _detectar_tesseract(log=None):
    """
    Procura o executável Tesseract em múltiplos locais possíveis.
    Retorna o caminho se encontrado, None caso contrário.
    """
    candidatos = [
        # Embutido no .exe (prioridade máxima)
        resource_path(os.path.join("tesseract", "tesseract.exe")),
        resource_path(os.path.join("assets", "tesseract", "tesseract.exe")),
        # Instalado pelo Chocolatey (GitHub Actions)
        r"C:\ProgramData\chocolatey\lib\tesseract\tools\tesseract.exe",
        r"C:\ProgramData\chocolatey\bin\tesseract.exe",
        # Instalação padrão UB-Mannheim
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        # PATH do sistema
        "tesseract",
    ]

    for caminho in candidatos:
        if caminho == "tesseract":
            # Tentar executar direto do PATH
            try:
                import subprocess
                subprocess.run(
                    ["tesseract", "--version"],
                    capture_output=True,
                    timeout=5,
                    creationflags=0x08000000,  # CREATE_NO_WINDOW
                )
                if log:
                    log("  ✓ Tesseract detectado no PATH do sistema")
                return "tesseract"
            except Exception:
                continue
        elif os.path.exists(caminho):
            if log:
                log(f"  ✓ Tesseract detectado em: {caminho}")
            return caminho

    if log:
        log("  ✗ Tesseract NÃO encontrado em nenhum local conhecido")
    return None


def _detectar_tessdata(log=None):
    """Localiza a pasta tessdata (idiomas do Tesseract)."""
    candidatos = [
        resource_path(os.path.join("tesseract", "tessdata")),
        resource_path(os.path.join("assets", "tesseract", "tessdata")),
        r"C:\ProgramData\chocolatey\lib\tesseract\tools\tessdata",
        r"C:\Program Files\Tesseract-OCR\tessdata",
        r"C:\Program Files (x86)\Tesseract-OCR\tessdata",
    ]
    for caminho in candidatos:
        if os.path.exists(caminho):
            # Verifica se tem ao menos o por.traineddata
            if os.path.exists(os.path.join(caminho, "por.traineddata")):
                if log:
                    log(f"  ✓ tessdata encontrado em: {caminho}")
                return caminho
    if log:
        log("  ⚠ tessdata 'por.traineddata' não encontrado — OCR pode falhar")
    return None


def _preprocessar_imagem_ocr(pil_img):
    """
    Pré-processa imagem para melhorar precisão do OCR:
    1. Converte para escala de cinza
    2. Upscale 3x com LANCZOS
    3. Aumenta contraste (autocontrast)
    4. Binariza (preto/branco puro)
    5. Aplica filtro de nitidez
    """
    img = pil_img.convert("L")  # escala de cinza
    # Upscale 3x
    img = img.resize((img.width * 3, img.height * 3), Image.LANCZOS)
    # Autocontraste
    img = ImageOps.autocontrast(img, cutoff=2)
    # Nitidez
    img = img.filter(ImageFilter.SHARPEN)
    # Binarização (threshold 160 funciona bem para documentos)
    img = img.point(lambda p: 255 if p > 160 else 0)
    return img


def ler_art_ocr(pdf_path, log=None):
    """
    Lê os campos da ART via OCR.
    Retorna dict com os campos extraídos (ou {} se falhar).
    """
    try:
        import pytesseract
        import fitz  # PyMuPDF
    except ImportError as e:
        if log:
            log(f"  ✗ Bibliotecas de OCR ausentes: {e}")
        return {}

    # Detectar Tesseract
    tess_cmd = _detectar_tesseract(log=log)
    if tess_cmd is None:
        if log:
            log("  ✗ Tesseract não instalado. Instale de: https://github.com/UB-Mannheim/tesseract/releases")
        return {}
    pytesseract.pytesseract.tesseract_cmd = tess_cmd

    # Detectar tessdata
    tessdata = _detectar_tessdata(log=log)
    cfg_tessdata = f'--tessdata-dir "{tessdata}"' if tessdata else ""

    # Renderizar PDF como imagem (300 DPI equivalente)
    try:
        doc = fitz.open(pdf_path)
        if len(doc) == 0:
            if log:
                log("  ✗ PDF vazio")
            return {}
        # Matrix(3.0, 3.0) = 216 DPI se PDF é 72 DPI
        pix = doc[0].get_pixmap(matrix=fitz.Matrix(3.0, 3.0))
        img = Image.open(BytesIO(pix.tobytes("png")))
        doc.close()
    except Exception as e:
        if log:
            log(f"  ✗ Erro ao renderizar PDF: {e}")
        return {}

    # Pré-processar imagem (MELHORIA v4)
    img = _preprocessar_imagem_ocr(img)

    # Executar OCR
    try:
        # PSM 6 = assume single uniform block of text (bom para ARTs)
        config = f'{cfg_tessdata} --psm 6 --oem 3'
        texto = pytesseract.image_to_string(img, lang="por+eng", config=config)
    except Exception as e:
        if log:
            log(f"  ✗ Erro no OCR: {e}")
        return {}

    if log:
        log(f"  • OCR extraiu {len(texto)} caracteres")

    return _extrair_campos_art(texto, log=log)


def _extrair_campos_art(texto, log=None):
    """
    Extrai campos da ART a partir do texto OCR.
    Regex TOLERANTES a variações de formato (v4).
    """
    campos = {}

    # Normaliza texto — remove múltiplos espaços
    texto_norm = re.sub(r"[ \t]+", " ", texto)

    # ART — 13 dígitos (pode ter espaços/pontos no meio)
    # Ex: "ART nº 1022722034987" ou "Nº da ART: 10 22 72 20 34 987"
    m = re.search(
        r"(?:ART|N[º°\.]?\s*(?:da\s*)?ART|N[ºº]mero)\s*[:\-]?\s*"
        r"([\d\.\s\-/]{10,20})",
        texto_norm, flags=re.IGNORECASE,
    )
    if m:
        num = re.sub(r"[^\d]", "", m.group(1))
        if 10 <= len(num) <= 15:
            campos["art"] = num

    # CREA — padrão NNNNNNNN[D/TD]-UF
    # Ex: "CREA 1022722034D-GO", "CREA:1022722034/D-GO"
    m = re.search(
        r"(\d{7,12}\s*[/\-]?\s*[A-Z]{1,3}\s*[/\-]?\s*[A-Z]{2})",
        texto_norm,
    )
    if m:
        crea = re.sub(r"\s+", "", m.group(1))
        campos["crea"] = crea

    # CEP — XXXXX-XXX ou XXXXXXXX
    m = re.search(r"(\d{5})\s*[\-\.]?\s*(\d{3})", texto_norm)
    if m:
        campos["cep"] = f"{m.group(1)}-{m.group(2)}"

    # Quadra / Lote
    # Ex: "Quadra 15 Lote 10", "QD 15 LT 10", "Qd. 15, Lt. 10"
    m = re.search(
        r"(?:quadra|qd\.?|Q)\s*[:\.]?\s*(\w{1,10})"
        r"[\s,\-e]+"
        r"(?:lote|lt\.?|L)\s*[:\.]?\s*(\w{1,10})",
        texto_norm, flags=re.IGNORECASE,
    )
    if m:
        campos["quadra_lote"] = f"QD {m.group(1).upper()} LT {m.group(2).upper()}"

    # CPF — XXX.XXX.XXX-XX
    m = re.search(r"(\d{3})\.?(\d{3})\.?(\d{3})\-?(\d{2})", texto_norm)
    if m:
        campos["cpf"] = f"{m.group(1)}.{m.group(2)}.{m.group(3)}-{m.group(4)}"

    # Cidade — procurar após "cidade" ou "município"
    m = re.search(
        r"(?:cidade|munic[ií]pio)\s*[:\-]?\s*([A-ZÀ-Ú][A-ZÀ-Úa-zà-ú\s]{2,30})",
        texto_norm, flags=re.IGNORECASE,
    )
    if m:
        campos["cidade"] = m.group(1).strip().upper()

    # Bairro
    m = re.search(
        r"(?:bairro|setor)\s*[:\-]?\s*([A-ZÀ-Ú][A-ZÀ-Úa-zà-ú\s]{2,30})",
        texto_norm, flags=re.IGNORECASE,
    )
    if m:
        campos["bairro"] = m.group(1).strip().upper()

    # Logradouro — procurar padrões "Rua X", "Av. Y", "Alameda Z"
    m = re.search(
        r"((?:rua|avenida|av\.|alameda|al\.|travessa|tv\.|rodovia|rod\.)\s+"
        r"[A-ZÀ-Ú][A-ZÀ-Úa-zà-ú\s\d]{3,50})",
        texto_norm, flags=re.IGNORECASE,
    )
    if m:
        campos["logradouro"] = m.group(1).strip().upper()

    if log:
        if campos:
            log(f"  ✓ Campos extraídos: {', '.join(campos.keys())}")
        else:
            log("  ⚠ Nenhum campo extraído — verifique qualidade do PDF")

    return campos



# ============================================================
# SCPO — FUNÇÕES DE AUTOMAÇÃO
# ============================================================

def _scpo_obter_chromedriver(log_cb=print):
    """Detecta versão do Chrome e baixa ChromeDriver compatível."""
    versao_major = "147"  # fallback — atualizar se o Chrome avançar muito
    try:
        for chave in [
            r"SOFTWARE\Google\Chrome\BLBeacon",
            r"SOFTWARE\Wow6432Node\Google\Chrome\BLBeacon",
        ]:
            try:
                with _winreg.OpenKey(_winreg.HKEY_LOCAL_MACHINE, chave) as k:
                    v, _ = _winreg.QueryValueEx(k, "version")
                    versao_major = v.split(".")[0]
                    log_cb(f"  Chrome detectado: v{v}")
                    break
            except Exception:
                pass
    except Exception:
        pass

    driver_dir  = Path.home() / "AppData" / "Local" / "SCPODriver"
    driver_path = driver_dir / "chromedriver.exe"
    versao_file = driver_dir / "versao.txt"
    driver_dir.mkdir(parents=True, exist_ok=True)

    if driver_path.exists() and versao_file.exists():
        cached = versao_file.read_text().strip()
        if cached.startswith(versao_major + "."):
            log_cb(f"  ChromeDriver {cached} em cache.")
            return str(driver_path)

    log_cb(f"  Baixando ChromeDriver para Chrome {versao_major}...")
    import ssl as _ssl
    _ctx = _ssl.create_default_context()
    _ctx.check_hostname = False
    _ctx.verify_mode = _ssl.CERT_NONE
    api = "https://googlechromelabs.github.io/chrome-for-testing/known-good-versions-with-downloads.json"
    with _urllib_scpo.urlopen(api, timeout=15, context=_ctx) as r:
        dados = _json_scpo.loads(r.read())

    url_zip = versao_exata = None
    for v in reversed(dados["versions"]):
        if v["version"].startswith(versao_major + "."):
            for d in v.get("downloads", {}).get("chromedriver", []):
                if d["platform"] == "win64":
                    url_zip = d["url"]
                    versao_exata = v["version"]
                    break
            if url_zip:
                break

    if not url_zip:
        raise Exception(f"ChromeDriver para Chrome {versao_major} não encontrado.")

    zip_path = driver_dir / "chromedriver.zip"
    opener = _urllib_scpo.build_opener(_urllib_scpo.HTTPSHandler(context=_ctx))
    with opener.open(url_zip) as resp, open(zip_path, "wb") as f:
        f.write(resp.read())
    with _zipfile_scpo.ZipFile(zip_path, "r") as z:
        for nome in z.namelist():
            if nome.endswith("chromedriver.exe"):
                with z.open(nome) as s, open(driver_path, "wb") as d:
                    d.write(s.read())
                break
    zip_path.unlink()
    versao_file.write_text(versao_exata)
    log_cb(f"  ChromeDriver {versao_exata} instalado.")
    return str(driver_path)


def _scpo_montar_nome_obra(logradouro, quadra, lote):
    return f"RESIDENCIAL {logradouro.upper()} QUADRA {quadra} LOTE {lote}"


def _scpo_montar_observacao(logradouro, quadra, lote, n_casas,
                             esquina, rua2, ruas_casas):
    casas = []
    for i in range(1, n_casas + 1):
        label = f"CASA {i}"
        if esquina and ruas_casas and i <= len(ruas_casas) and ruas_casas[i-1]:
            label += f" SITUADA NA {ruas_casas[i-1].upper()}"
        casas.append(label)
    casas_str = ", ".join(casas)
    if not esquina:
        return (f"OBRA RESIDENCIAL UNIFAMILIAR SITUADA NA {logradouro.upper()}, "
                f"QUADRA {quadra} LOTE {lote} COMPOSTA POR: {casas_str}")
    return (f"OBRA RESIDENCIAL UNIFAMILIAR SITUADA NA {logradouro.upper()} "
            f"E {rua2.upper()}, QUADRA {quadra} LOTE {lote} "
            f"COMPOSTA POR: {casas_str}")


def _scpo_data_termino(data_inicio_str):
    """Retorna data de término = 1 mês após data_inicio_str (DD/MM/AAAA)."""
    from dateutil.relativedelta import relativedelta
    dt = datetime.datetime.strptime(data_inicio_str, "%d/%m/%Y")
    return (dt + relativedelta(months=1)).strftime("%d/%m/%Y")


def _scpo_executar(dados, step_cb, log_cb, done_cb,
                   evento_captcha, fn_habilitar_captcha,
                   evento_envio, fn_habilitar_envio):
    """Thread principal de automação SCPO."""
    import time, traceback

    LOGIN_CPF     = "038.144.411-25"
    EMAIL_FIXO    = "joaovitorcabral94@gmail.com"
    TELEFONE_FIXO = "(62)99266-5923"
    EMP_PRINCIPAL = "0"
    EMP_TERCEIROS = "5"
    URL_LOGIN     = "https://scpo.mte.gov.br/"

    driver = None
    try:
        step_cb(5, "Obtendo ChromeDriver...")
        driver_path = _scpo_obter_chromedriver(log_cb)

        options = _webdriver_scpo.ChromeOptions()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-notifications")
        driver = _webdriver_scpo.Chrome(
            service=_Service_scpo(driver_path), options=options)
        wait = _Wait_scpo(driver, 30)

        # LOGIN
        step_cb(10, "Abrindo SCPO...")
        log_cb("Abrindo SCPO...")
        driver.get(URL_LOGIN)
        wait.until(_EC_scpo.presence_of_element_located((_By_scpo.ID, "txtCPF")))
        driver.find_element(_By_scpo.ID, "txtCPF").send_keys(LOGIN_CPF)
        driver.find_element(_By_scpo.ID, "PlaceHolderConteudo_txtSenha").send_keys(dados["senha"])
        log_cb("CPF e senha preenchidos. Digite o captcha no navegador.")
        step_cb(15, "Aguardando captcha...")
        fn_habilitar_captcha()
        evento_captcha.wait()
        log_cb("Captcha confirmado. Entrando...")
        step_cb(20, "Efetuando login...")
        driver.find_element(_By_scpo.ID, "PlaceHolderConteudo_btnLogin").click()
        wait.until(_EC_scpo.presence_of_element_located(
            (_By_scpo.XPATH, "//a[contains(@onclick,'subMenu01')]")))
        log_cb("✓ Login OK!")

        # NAVEGAÇÃO
        step_cb(30, "Navegando para Comunicar Obra...")
        wait.until(_EC_scpo.element_to_be_clickable(
            (_By_scpo.XPATH, "//a[contains(@onclick,'subMenu01')]"))).click()
        time.sleep(1)
        wait.until(_EC_scpo.element_to_be_clickable(
            (_By_scpo.XPATH, "//a[contains(@href,'DeclaracaoPreviaObra/Comunicar')]"))).click()

        # TELA INTERMEDIÁRIA
        step_cb(38, "Identificando empresa...")
        chk = wait.until(_EC_scpo.presence_of_element_located(
            (_By_scpo.ID, "PlaceHolderConteudo_chkObraSemCNPJ")))
        if not chk.is_selected():
            chk.click()
        time.sleep(1)
        cpf_f = wait.until(_EC_scpo.element_to_be_clickable(
            (_By_scpo.ID, "txtCPFProprietarioObra")))
        cpf_f.clear()
        cpf_f.send_keys(LOGIN_CPF)
        driver.find_element(
            _By_scpo.ID, "PlaceHolderConteudo_btnDeclararObra").click()

        # FORMULÁRIO
        step_cb(50, "Preenchendo formulário...")
        log_cb("Preenchendo formulário...")
        wait.until(_EC_scpo.presence_of_element_located(
            (_By_scpo.ID, "txtNomeObraEmpreendimento"))).send_keys(dados["nome_obra"])
        log_cb(f"  Nome: {dados['nome_obra']}")
        driver.find_element(_By_scpo.ID, "txtEmailObra").send_keys(EMAIL_FIXO)
        driver.find_element(_By_scpo.ID, "txtTelefoneObra").send_keys(TELEFONE_FIXO)

        # CEP dígito por dígito
        cep_limpo = re.sub(r"[^0-9]", "", dados["cep"])
        campo_cep = driver.find_element(_By_scpo.ID, "txtObraCEP")
        driver.execute_script("arguments[0].value = '';", campo_cep)
        driver.execute_script("arguments[0].focus();", campo_cep)
        for digito in cep_limpo:
            campo_cep.send_keys(digito)
            time.sleep(0.05)
        time.sleep(1)
        driver.find_element(
            _By_scpo.ID, "PlaceHolderConteudo_imgPesquisarCEPObra").click()
        time.sleep(4)
        log_cb(f"  CEP: {cep_limpo}")

        # Número: SN
        try:
            n = driver.find_element(_By_scpo.ID, "txtObraNumero")
            n.clear(); n.send_keys("SN")
        except Exception: pass

        # Complemento
        try:
            c = driver.find_element(_By_scpo.XPATH,
                "//input[contains(@id,'Complemento') or contains(@id,'complemento')]")
            c.clear()
            c.send_keys(f"QUADRA {dados['quadra']} LOTE {dados['lote']}")
        except Exception: pass

        # Observação
        try:
            obs = driver.find_element(_By_scpo.XPATH,
                "//textarea[contains(@id,'Observ') or contains(@id,'observ') "
                "or contains(@id,'Descri')]")
            obs.clear()
            obs.send_keys(dados["observacao"])
            log_cb(f"  Obs: {dados['observacao'][:80]}...")
        except Exception as e:
            log_cb(f"  textarea não localizada: {e}")

        # Classe CNAE 4120-4
        step_cb(65, "Selecionando CNAE e tipo...")
        try:
            sel = _Select_scpo(driver.find_element(
                _By_scpo.ID, "PlaceHolderConteudo_cboClasseCNAE"))
            for opt in sel.options:
                if "4120" in opt.text:
                    sel.select_by_visible_text(opt.text); break
            time.sleep(1)
        except Exception as e: log_cb(f"  CNAE: {e}")

        # Subclasse 00
        try:
            sel_sub = _Select_scpo(wait.until(_EC_scpo.presence_of_element_located(
                (_By_scpo.ID, "PlaceHolderConteudo_cboSubclasse"))))
            for opt in sel_sub.options:
                if opt.text.strip().startswith("00"):
                    sel_sub.select_by_visible_text(opt.text); break
        except Exception as e: log_cb(f"  Subclasse: {e}")

        # Tipo Construção — Edifício
        try:
            sel_tipo = _Select_scpo(driver.find_element(
                _By_scpo.ID, "PlaceHolderConteudo_CboTipoConstrucao"))
            for opt in sel_tipo.options:
                if "dif" in opt.text.lower():
                    sel_tipo.select_by_visible_text(opt.text); break
        except Exception as e: log_cb(f"  Tipo Construção: {e}")

        # Tipo Obra — Privada
        try:
            driver.find_element(_By_scpo.XPATH,
                "//input[@type='radio' and contains(@id,'rivada')]").click()
        except Exception as e: log_cb(f"  Privada: {e}")

        # Característica — Construção
        try:
            driver.find_element(_By_scpo.XPATH,
                "//input[@type='radio' and "
                "(@value='Construcao' or @value='Construção' "
                "or contains(@id,'onstrucao'))]").click()
        except Exception as e: log_cb(f"  Construção: {e}")

        # FGTS — Não
        try:
            driver.find_element(
                _By_scpo.ID, "PlaceHolderConteudo_rdbFinanciamentoFGTSNao").click()
        except Exception as e: log_cb(f"  FGTS: {e}")

        # Datas
        try:
            from dateutil.relativedelta import relativedelta
            data_termino = (datetime.date.today() + relativedelta(months=1)).strftime("%d/%m/%Y")
            driver.find_element(_By_scpo.ID, "txtInicio").send_keys(dados["data_inicio"])
            driver.find_element(_By_scpo.ID, "txtTermino").send_keys(data_termino)
            log_cb(f"  Datas: {dados['data_inicio']} → {data_termino}")
        except Exception as e: log_cb(f"  Datas: {e}")

        # Empregados
        try:
            driver.find_element(
                _By_scpo.ID, "txtNumeroEmpregadosEmpresaPrincipal").send_keys(EMP_PRINCIPAL)
        except Exception as e: log_cb(f"  Emp. principal: {e}")
        try:
            driver.find_element(_By_scpo.XPATH,
                "//input[contains(@id,'Terceiros') or contains(@id,'terceiros')]"
            ).send_keys(EMP_TERCEIROS)
        except Exception as e: log_cb(f"  Emp. terceiros: {e}")

        step_cb(85, "Aguardando confirmação...")
        log_cb("✓ Formulário preenchido! Verifique o navegador.")
        log_cb("Clique 'Confirmar envio SCPO' quando estiver pronto.")
        fn_habilitar_envio()
        evento_envio.wait()

        # SUBMETER
        step_cb(92, "Submetendo...")
        log_cb("Submetendo formulário...")
        try:
            driver.find_element(
                _By_scpo.ID, "PlaceHolderConteudo_btnConfirmar").click()
            time.sleep(3)
            log_cb("✓ Formulário submetido!")
        except Exception as e:
            log_cb(f"  Erro ao submeter: {e}")

        step_cb(100, "SCPO concluído!")
        done_cb(True, "SCPO preenchido com sucesso!")

    except Exception as e:
        tb = traceback.format_exc()
        log_cb(f"✗ ERRO SCPO: {e}")
        log_cb(tb)
        done_cb(False, str(e))


# ============================================================
# INTERFACE TKINTER
# ============================================================

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("BERÇAN PROJETOS — Preenchimento de Documentos")
        self.geometry("1100x780")
        self.configure(bg=COR_FUNDO)

        # Controle de interrupção (NOVO v4)
        self.stop_event = threading.Event()
        self.processando = False

        self._criar_widgets()

        # Restaurar configurações salvas
        cfg = _config_carregar()
        if cfg.get("scpo_senha"):
            self.var_scpo_senha.set(cfg["scpo_senha"])

    # ------------------------------------------------------------------
    # Construção da UI
    # ------------------------------------------------------------------
    def _criar_widgets(self):
        # Cabeçalho
        header = tk.Frame(self, bg=COR_FUNDO)
        header.pack(fill="x", padx=20, pady=(20, 10))

        # Logo + título lado a lado
        header_inner = tk.Frame(header, bg=COR_FUNDO)
        header_inner.pack(anchor="w")

        # Logo (carrega LOGO.png dos assets)
        self._logo_img = None
        try:
            from PIL import Image, ImageTk
            logo_path = asset("LOGO.png")
            if not os.path.exists(logo_path):
                logo_path = asset("LOGO.jpg")
            if os.path.exists(logo_path):
                img = Image.open(logo_path).convert("RGBA")
                img = img.resize((52, 52), Image.LANCZOS)
                self._logo_img = ImageTk.PhotoImage(img)
                tk.Label(header_inner, image=self._logo_img,
                         bg=COR_FUNDO).pack(side="left", padx=(0, 12))
        except Exception:
            pass  # sem logo, continua normal

        texto_frame = tk.Frame(header_inner, bg=COR_FUNDO)
        texto_frame.pack(side="left")
        tk.Label(
            texto_frame, text="BERÇAN PROJETOS",
            font=("Segoe UI", 18, "bold"),
            fg=COR_TEXTO, bg=COR_FUNDO,
        ).pack(anchor="w")
        tk.Label(
            texto_frame, text="Preenchimento Automático de Documentos",
            font=("Segoe UI", 10),
            fg=COR_TEXTO_SEC, bg=COR_FUNDO,
        ).pack(anchor="w")

        # Container principal em 2 colunas
        main = tk.Frame(self, bg=COR_FUNDO)
        main.pack(fill="both", expand=True, padx=20, pady=10)

        col_esq = tk.Frame(main, bg=COR_FUNDO)
        col_esq.pack(side="left", fill="both", expand=True, padx=(0, 10))

        col_dir = tk.Frame(main, bg=COR_FUNDO)
        col_dir.pack(side="right", fill="both", expand=True, padx=(10, 0))

        # --- Coluna esquerda ---
        self._secao_label(col_esq, "MEMORIAL EXCEL")
        self.var_memorial = tk.StringVar()
        self._campo_arquivo(col_esq, self.var_memorial, "Arquivo Memorial (.xls/.xlsx)")

        self._secao_label(col_esq, "ENGENHEIRO RESPONSÁVEL")
        self.var_engenheiro = tk.StringVar()
        combo_eng = ttk.Combobox(
            col_esq, textvariable=self.var_engenheiro,
            values=list(ENGENHEIROS.keys()),
            state="readonly", font=("Segoe UI", 10),
        )
        combo_eng.pack(fill="x", pady=3)
        combo_eng.bind("<<ComboboxSelected>>", self._preencher_eng_campos)

        self.var_cpf = tk.StringVar()
        self.var_crea = tk.StringVar()
        self._campo_simples(col_esq, self.var_cpf, "CPF")
        self._campo_simples(col_esq, self.var_crea, "CREA")

        self._secao_label(col_esq, "DADOS DA ART")

        self.var_art = tk.StringVar()
        self._campo_simples(col_esq, self.var_art, "Número da ART")
        self.var_registro_crea = tk.StringVar()
        self._campo_simples(col_esq, self.var_registro_crea, "Registro CREA")

        self.var_contratante = tk.StringVar()
        self._campo_simples(col_esq, self.var_contratante, "Contratante")

        # CEP com botão de busca automática — PRIMEIRO para auto-preencher
        self.var_cep = tk.StringVar()
        # CEP com botão de busca automática
        tk.Label(col_esq, text="CEP",
                 bg=COR_FUNDO, fg=COR_TEXTO_SEC, font=("Segoe UI", 8)).pack(anchor="w")
        frame_cep = tk.Frame(col_esq, bg=COR_FUNDO)
        frame_cep.pack(fill="x", pady=(0, 3))
        tk.Entry(frame_cep, textvariable=self.var_cep,
                 bg=COR_CAMPO, fg=COR_TEXTO, insertbackground=COR_TEXTO,
                 relief="flat", font=("Segoe UI", 10),
                 width=12).pack(side="left")
        tk.Button(frame_cep, text="🔍 Buscar CEP",
                  command=self._buscar_cep,
                  bg=COR_BOTAO, fg=COR_TEXTO, relief="flat",
                  font=("Segoe UI", 9, "bold"),
                  ).pack(side="left", padx=(6, 0))
        self.lbl_cep_status = tk.Label(frame_cep, text="",
                  bg=COR_FUNDO, fg=COR_LOG_TEXTO,
                  font=("Segoe UI", 8))
        self.lbl_cep_status.pack(side="left", padx=(6, 0))

        self.var_logradouro = tk.StringVar()
        self._campo_simples(col_esq, self.var_logradouro, "Logradouro")

        self.var_quadra_lote = tk.StringVar()
        self._campo_simples(col_esq, self.var_quadra_lote, "Quadra e Lote")

        self.var_bairro = tk.StringVar()
        self._campo_simples(col_esq, self.var_bairro, "Bairro")

        self.var_cidade = tk.StringVar()
        self._campo_simples(col_esq, self.var_cidade, "Cidade")

        self.var_uf = tk.StringVar(value="GO")
        self._campo_simples(col_esq, self.var_uf, "UF")

        # ── SCPO: campos movidos para coluna direita (scroll_frame) ──────────
        self.var_scpo_data_inicio = tk.StringVar(
            value=datetime.date.today().strftime("%d/%m/%Y")
        )
        self.var_scpo_senha = tk.StringVar()

        # --- Coluna direita com Canvas+Scroll para todo o conteúdo ---
        # Botões fixos no rodapé (fora do scroll)
        frame_botoes = tk.Frame(col_dir, bg=COR_FUNDO)
        frame_botoes.pack(side="bottom", fill="x", pady=(6,0))
        self.btn_gerar = tk.Button(
            frame_botoes, text="⚡ GERAR DOCUMENTOS",
            command=self._iniciar_geracao,
            bg=COR_BOTAO, fg=COR_TEXTO, relief="flat",
            font=("Segoe UI", 12, "bold"), padx=20, pady=10,
        )
        self.btn_gerar.pack(side="left", fill="x", expand=True, padx=(0, 5))
        self.btn_stop = tk.Button(
            frame_botoes, text="⛔ INTERROMPER",
            command=self._solicitar_stop,
            bg=COR_BOTAO_STOP, fg=COR_TEXTO, relief="flat",
            font=("Segoe UI", 12, "bold"), padx=20, pady=10,
            state="disabled",
        )
        self.btn_stop.pack(side="right", fill="x", expand=True, padx=(5, 0))

        # Botão Calibrar — linha abaixo de GERAR/INTERROMPER
        frame_calibrar = tk.Frame(col_dir, bg=COR_FUNDO)
        frame_calibrar.pack(side="bottom", fill="x", pady=(0, 4))
        tk.Button(
            frame_calibrar, text="⚙ CALIBRAR MEMORIAL",
            command=self._abrir_calibrador,
            bg="#4a4a6a", fg=COR_TEXTO, relief="flat",
            font=("Segoe UI", 10, "bold"), pady=6,
        ).pack(fill="x")

        # Botão SCPO — linha separada abaixo
        frame_botoes2 = tk.Frame(col_dir, bg=COR_FUNDO)
        frame_botoes2.pack(side="bottom", fill="x", pady=(4, 0))
        self.btn_scpo = tk.Button(
            frame_botoes2, text="🌐 PREENCHER SCPO",
            command=self._iniciar_scpo,
            bg="#1a6b3c", fg=COR_TEXTO, relief="flat",
            font=("Segoe UI", 11, "bold"), padx=20, pady=8,
        )
        self.btn_scpo.pack(side="left", fill="x", expand=True, padx=(0, 5))
        self.btn_scpo_captcha = tk.Button(
            frame_botoes2, text="✔ Código digitado",
            command=self._scpo_liberar_captcha,
            bg="#27ae60", fg=COR_TEXTO, relief="flat",
            font=("Segoe UI", 10), padx=10, pady=8,
            state="disabled",
        )
        self.btn_scpo_captcha.pack(side="left", padx=(0, 5))
        self.btn_scpo_envio = tk.Button(
            frame_botoes2, text="✔ Confirmar envio SCPO",
            command=self._scpo_liberar_envio,
            bg="#e67e22", fg=COR_TEXTO, relief="flat",
            font=("Segoe UI", 10), padx=10, pady=8,
            state="disabled",
        )
        self.btn_scpo_envio.pack(side="left")

        # Canvas scrollável para o restante da coluna direita
        canvas_dir = tk.Canvas(col_dir, bg=COR_FUNDO, highlightthickness=0)
        sb_dir = tk.Scrollbar(col_dir, orient="vertical", command=canvas_dir.yview)
        canvas_dir.configure(yscrollcommand=sb_dir.set)
        sb_dir.pack(side="right", fill="y")
        canvas_dir.pack(side="left", fill="both", expand=True)
        scroll_frame = tk.Frame(canvas_dir, bg=COR_FUNDO)
        wid_dir = canvas_dir.create_window((0, 0), window=scroll_frame, anchor="nw")
        def _on_cf(e):
            canvas_dir.configure(scrollregion=canvas_dir.bbox("all"))
            canvas_dir.itemconfig(wid_dir, width=canvas_dir.winfo_width())
        scroll_frame.bind("<Configure>", _on_cf)
        # Forçar render inicial
        self.after(50, lambda: canvas_dir.configure(
            scrollregion=canvas_dir.bbox("all")))
        self.after(100, lambda: canvas_dir.itemconfig(
            wid_dir, width=canvas_dir.winfo_width()))
        canvas_dir.bind("<MouseWheel>",
            lambda e: canvas_dir.yview_scroll(int(-1*(e.delta/120)), "units"))
        scroll_frame.bind("<MouseWheel>",
            lambda e: canvas_dir.yview_scroll(int(-1*(e.delta/120)), "units"))
        p = scroll_frame  # alias — todo widget da col_dir usa p

        # ── 1. QUANTIDADE DE CASAS ──
        self._secao_label(p, "QUANTIDADE DE CASAS")
        self.var_qtd_casas = tk.IntVar(value=1)
        tk.Spinbox(
            p, from_=1, to=50, textvariable=self.var_qtd_casas,
            width=5, bg=COR_CAMPO, fg=COR_TEXTO,
            insertbackground=COR_TEXTO, relief="flat",
        ).pack(anchor="w", pady=3)

        # ── 2. LOTE DE ESQUINA ──
        self._secao_label(p, "LOTE")
        self.var_esquina = tk.BooleanVar(value=False)
        self._chk_esquina_widget = tk.Checkbutton(
            p, text="Lote de esquina (frente para mais de uma rua)",
            variable=self.var_esquina,
            bg=COR_FUNDO, fg=COR_TEXTO, selectcolor=COR_CAMPO,
            activebackground=COR_FUNDO, activeforeground=COR_TEXTO,
            command=self._toggle_ruas_esquina,
        )
        self._chk_esquina_widget.pack(anchor="w", pady=3)
        # Frame de ruas — campos dinâmicos (um por casa)
        self.frame_ruas_esquina = tk.Frame(p, bg=COR_FUNDO)
        self._canvas_dir = canvas_dir  # referência para scroll
        self._p_dir = p                # referência para criar campos
        self._entries_ruas = []        # lista de Entry, um por casa
        # Tracer: recria campos quando qtd_casas muda
        self.var_qtd_casas.trace_add("write",
            lambda *_: self.after(100, self._rebuild_ruas_esquina))

        # ── 3. OPÇÕES ──
        self._secao_label(p, "OPÇÕES")
        self.var_esgoto = tk.BooleanVar(value=False)
        tk.Checkbutton(
            p, text="Sistema público de esgoto (SIM)",
            variable=self.var_esgoto,
            bg=COR_FUNDO, fg=COR_TEXTO, selectcolor=COR_CAMPO,
            activebackground=COR_FUNDO, activeforeground=COR_TEXTO,
        ).pack(anchor="w", pady=3)
        tk.Label(
            p, text="↳ Template Word selecionado automaticamente",
            bg=COR_FUNDO, fg=COR_TEXTO_SEC, font=("Segoe UI", 8),
        ).pack(anchor="w", padx=20)

        self._secao_label(p, "CASAS GEMINADAS")
        _og = ["Não se aplica", "Sim", "Não"]
        tk.Label(p, text="Condomínios",
                 bg=COR_FUNDO, fg=COR_TEXTO_SEC, font=("Segoe UI", 8)).pack(anchor="w")
        self.var_gem_cond = tk.StringVar(value="Não se aplica")
        ttk.Combobox(p, textvariable=self.var_gem_cond, values=_og,
                     state="readonly", font=("Segoe UI", 9)).pack(fill="x", pady=(0, 8))

        # ── SCPO ──
        self._secao_label(p, "SCPO")
        self._campo_simples(p, self.var_scpo_data_inicio,
                            "Data de Início da Obra (DD/MM/AAAA)")
        tk.Label(p, text="Senha SCPO",
                 bg=COR_FUNDO, fg=COR_TEXTO_SEC, font=("Segoe UI", 8)
                 ).pack(anchor="w")
        frame_senha_scpo = tk.Frame(p, bg=COR_FUNDO)
        frame_senha_scpo.pack(fill="x", pady=(0, 4))
        self._ent_scpo_senha = tk.Entry(
            frame_senha_scpo, textvariable=self.var_scpo_senha,
            bg=COR_CAMPO, fg=COR_TEXTO, insertbackground=COR_TEXTO,
            relief="flat", font=("Segoe UI", 10), show="*"
        )
        self._ent_scpo_senha.pack(side="left", fill="x", expand=True)
        tk.Button(
            frame_senha_scpo, text="👁",
            bg=COR_CAMPO, fg=COR_TEXTO, relief="flat",
            command=lambda: self._ent_scpo_senha.config(
                show="" if self._ent_scpo_senha.cget("show") == "*" else "*")
        ).pack(side="left", padx=(4, 0))
        # Salvar senha automaticamente sempre que for alterada
        self.var_scpo_senha.trace_add(
            "write",
            lambda *_: _config_salvar({"scpo_senha": self.var_scpo_senha.get()})
        )

        # ── LOG ──
        self._secao_label(p, "LOG")
        frame_log = tk.Frame(p, bg=COR_LOG_FUNDO)
        frame_log.pack(fill="x", pady=3)
        sb_log = tk.Scrollbar(frame_log, orient="vertical")
        sb_log.pack(side="right", fill="y")
        self.txt_log = tk.Text(
            frame_log, height=10, bg=COR_LOG_FUNDO, fg=COR_LOG_TEXTO,
            font=("Consolas", 9), relief="flat",
            yscrollcommand=sb_log.set,
        )
        self.txt_log.pack(side="left", fill="both", expand=True)
        sb_log.config(command=self.txt_log.yview)

        # ── PROGRESSO ──
        self._secao_label(p, "PROGRESSO")
        self.progress = ttk.Progressbar(p, mode="determinate", length=400)
        self.progress.pack(fill="x", pady=3)
        self.var_status = tk.StringVar(value="Aguardando...")
        tk.Label(
            p, textvariable=self.var_status,
            bg=COR_FUNDO, fg=COR_TEXTO_SEC, font=("Segoe UI", 9),
        ).pack(anchor="w", pady=(0, 10))

    # ------------------------------------------------------------------
    # Helpers de UI
    # ------------------------------------------------------------------
    def _secao_label(self, parent, texto):
        tk.Label(
            parent, text=texto,
            bg=COR_FUNDO, fg=COR_TEXTO, font=("Segoe UI", 10, "bold"),
        ).pack(anchor="w", pady=(10, 3))

    def _campo_simples(self, parent, var, hint):
        tk.Label(
            parent, text=hint,
            bg=COR_FUNDO, fg=COR_TEXTO_SEC, font=("Segoe UI", 8),
        ).pack(anchor="w")
        tk.Entry(
            parent, textvariable=var,
            bg=COR_CAMPO, fg=COR_TEXTO, insertbackground=COR_TEXTO,
            relief="flat", font=("Segoe UI", 10),
        ).pack(fill="x", pady=(0, 3))

    def _campo_arquivo(self, parent, var, hint):
        tk.Label(
            parent, text=hint,
            bg=COR_FUNDO, fg=COR_TEXTO_SEC, font=("Segoe UI", 8),
        ).pack(anchor="w")
        frame = tk.Frame(parent, bg=COR_FUNDO)
        frame.pack(fill="x", pady=(0, 3))
        tk.Entry(
            frame, textvariable=var,
            bg=COR_CAMPO, fg=COR_TEXTO, insertbackground=COR_TEXTO,
            relief="flat", font=("Segoe UI", 10),
        ).pack(side="left", fill="x", expand=True)
        tk.Button(
            frame, text="📁", command=lambda: self._selecionar_arquivo(
                var, [("Excel", "*.xls *.xlsx")]),
            bg=COR_BOTAO, fg=COR_TEXTO, relief="flat", font=("Segoe UI", 9),
        ).pack(side="right", padx=(5, 0))

    def _selecionar_arquivo(self, var, filetypes):
        caminho = filedialog.askopenfilename(filetypes=filetypes)
        if caminho:
            var.set(caminho)

    def _preencher_eng_campos(self, _event=None):
        nome = self.var_engenheiro.get()
        if nome in ENGENHEIROS:
            self.var_cpf.set(ENGENHEIROS[nome]["cpf"])
            self.var_crea.set(ENGENHEIROS[nome]["crea"])
            self.var_registro_crea.set(ENGENHEIROS[nome]["crea"])

    def log(self, msg):
        """Adiciona mensagem ao log em tempo real (thread-safe via after)."""
        self.after(0, self._log_insert, msg)

    def _log_insert(self, msg):
        self.txt_log.insert("end", msg + "\n")
        self.txt_log.see("end")

    def _set_status(self, texto):
        self.after(0, self.var_status.set, texto)

    def _set_progress(self, valor):
        self.after(0, lambda: self.progress.configure(value=valor))

    # ------------------------------------------------------------------
    # Ações
    # ------------------------------------------------------------------
    def _acionar_ocr(self):
        """Dispara leitura OCR da ART em thread separada."""
        pdf = self.var_art_pdf.get().strip()
        if not pdf or not os.path.exists(pdf):
            messagebox.showerror("ART", "Selecione um PDF de ART válido.")
            return
        self.log("🔍 Iniciando OCR da ART...")
        threading.Thread(target=self._ocr_worker, args=(pdf,), daemon=True).start()

    def _ocr_worker(self, pdf):
        try:
            campos = ler_art_ocr(pdf, log=self.log)
            if not campos:
                self.log("✗ OCR não retornou dados utilizáveis")
                return
            # Preencher campos na UI (thread-safe)
            if "art" in campos:
                self.after(0, self.var_art.set, campos["art"])
            if "crea" in campos:
                self.after(0, self.var_registro_crea.set, campos["crea"])
            if "cep" in campos:
                self.after(0, self.var_cep.set, campos["cep"])
            if "quadra_lote" in campos:
                self.after(0, self.var_quadra_lote.set, campos["quadra_lote"])
            if "cidade" in campos:
                self.after(0, self.var_cidade.set, campos["cidade"])
            if "bairro" in campos:
                self.after(0, self.var_bairro.set, campos["bairro"])
            if "logradouro" in campos:
                self.after(0, self.var_logradouro.set, campos["logradouro"])
            self.log(f"✓ OCR concluído — {len(campos)} campos preenchidos")
        except Exception as e:
            self.log(f"✗ Erro OCR: {e}")
            self.log(traceback.format_exc())

    def _solicitar_stop(self):
        """Sinaliza para a thread de processamento parar (NOVO v4)."""
        if not self.processando:
            return
        self.stop_event.set()
        self.log("⚠️ Interrupção solicitada — aguardando etapa atual finalizar...")
        self.btn_stop.configure(state="disabled", text="⏳ PARANDO...")

    def _iniciar_geracao(self):
        """Valida entradas e dispara thread de processamento."""
        if not self.var_memorial.get() or not os.path.exists(self.var_memorial.get()):
            messagebox.showerror("Erro", "Selecione um arquivo Memorial Excel válido.")
            return
        if not self.var_engenheiro.get():
            messagebox.showerror("Erro", "Selecione o engenheiro responsável.")
            return
        if not self.var_art.get().strip():
            messagebox.showerror("Erro", "Informe o número da ART.")
            return

        # Reset
        self.stop_event.clear()
        self.processando = True
        self.btn_gerar.configure(state="disabled")
        self.btn_stop.configure(state="normal", text="⛔ INTERROMPER")
        self.progress.configure(value=0)
        self.txt_log.delete("1.0", "end")

        threading.Thread(target=self._processar, daemon=True).start()

    def _abrir_calibrador(self):
        """Abre a janela do Calibrador do Memorial."""
        memorial = self.var_memorial.get().strip()
        JanelaCalibrador(self, memorial if os.path.exists(memorial) else None)

    def _buscar_cep(self):
        """Consulta ViaCEP e preenche logradouro, bairro, cidade e UF."""
        cep = self.var_cep.get().strip()
        self.lbl_cep_status.configure(text="⏳ buscando...", fg=COR_TEXTO_SEC)

        def ok(data):
            self.after(0, self.var_logradouro.set,
                       data.get("logradouro", "").upper())
            self.after(0, self.var_bairro.set,
                       data.get("bairro", "").upper())
            self.after(0, self.var_cidade.set,
                       data.get("localidade", "").upper())
            self.after(0, self.var_uf.set,
                       data.get("uf", "GO").upper())
            self.after(0, self.lbl_cep_status.configure,
                       {"text": "✓ preenchido!", "fg": COR_LOG_TEXTO})

        def erro(msg):
            self.after(0, self.lbl_cep_status.configure,
                       {"text": f"✗ {msg}", "fg": "#e74c3c"})

        buscar_cep(cep, ok, erro)

    def _toggle_ruas_esquina(self):
        """Mostra/oculta e reconstrói campos de rua por casa."""
        if self.var_esquina.get():
            self._rebuild_ruas_esquina()
            self.frame_ruas_esquina.pack(fill="x", pady=(0, 6),
                                         after=self._chk_esquina_widget)
        else:
            self.frame_ruas_esquina.pack_forget()

    def _rebuild_ruas_esquina(self):
        """Reconstrói um Entry por casa dentro do frame_ruas_esquina."""
        if not self.var_esquina.get():
            return
        # Limpar campos anteriores
        for w in self.frame_ruas_esquina.winfo_children():
            w.destroy()
        self._entries_ruas = []
        try:
            qtd = int(self.var_qtd_casas.get())
        except:
            qtd = 1
        for i in range(1, qtd + 1):
            tk.Label(self.frame_ruas_esquina,
                     text=f"Rua — Casa {i}:",
                     bg=COR_FUNDO, fg=COR_TEXTO_SEC,
                     font=("Segoe UI", 8)).pack(anchor="w")
            var_rua = tk.StringVar()
            e = tk.Entry(self.frame_ruas_esquina,
                         textvariable=var_rua,
                         bg=COR_CAMPO, fg=COR_TEXTO,
                         insertbackground=COR_TEXTO,
                         relief="flat", font=("Segoe UI", 10))
            e.pack(fill="x", pady=(0, 4))
            e.bind("<MouseWheel>", lambda ev:
                   self._canvas_dir.yview_scroll(
                       int(-1*(ev.delta/120)), "units"))
            self._entries_ruas.append(var_rua)
        # Forçar atualização do canvas
        self.after(50, lambda: self._canvas_dir.configure(
            scrollregion=self._canvas_dir.bbox("all")))

    def _get_rua_casa(self, num_casa):
        """Retorna a rua da casa N (campo individual) ou logradouro padrão."""
        if not self.var_esquina.get():
            return self.var_logradouro.get()
        idx = num_casa - 1
        if idx < len(self._entries_ruas):
            val = self._entries_ruas[idx].get().strip()
            return val if val else self.var_logradouro.get()
        return self.var_logradouro.get()

    # ── Métodos SCPO ─────────────────────────────────────────────────────────
    def _scpo_liberar_captcha(self):
        self._scpo_evento_captcha.set()
        self.after(0, lambda: self.btn_scpo_captcha.config(state="disabled"))
        self.log("  ✓ Captcha SCPO confirmado.")

    def _scpo_liberar_envio(self):
        self._scpo_evento_envio.set()
        self.after(0, lambda: self.btn_scpo_envio.config(state="disabled"))
        self.log("  ✓ Envio SCPO confirmado.")

    def _iniciar_scpo(self):
        """Valida campos e dispara automação SCPO em thread separada."""
        # Validações
        if not self.var_cep.get().strip():
            messagebox.showwarning("SCPO", "Preencha o CEP antes de iniciar o SCPO.")
            return
        if not self.var_logradouro.get().strip():
            messagebox.showwarning("SCPO", "Preencha o Logradouro antes de iniciar o SCPO.")
            return
        ql = self.var_quadra_lote.get().strip()
        if not ql:
            messagebox.showwarning("SCPO", "Preencha Quadra e Lote antes de iniciar o SCPO.")
            return
        if not self.var_scpo_data_inicio.get().strip():
            messagebox.showwarning("SCPO", "Preencha a Data de Início da Obra (DD/MM/AAAA).")
            return
        if not self.var_scpo_senha.get().strip():
            messagebox.showwarning("SCPO", "Preencha a Senha SCPO.")
            return
        try:
            datetime.datetime.strptime(self.var_scpo_data_inicio.get().strip(), "%d/%m/%Y")
        except ValueError:
            messagebox.showwarning("SCPO", "Data de Início inválida. Use DD/MM/AAAA.")
            return

        # Separar quadra e lote de "QUADRA X LOTE Y"
        import re as _re
        ql_upper = ql.upper()
        m_q = _re.search(r"QUADRA\s+(\S+)", ql_upper)
        m_l = _re.search(r"LOTE\s+(\S+)", ql_upper)
        quadra = m_q.group(1) if m_q else ql
        lote   = m_l.group(1) if m_l else ""

        # Ruas por casa (esquina)
        n_casas   = self.var_qtd_casas.get()
        ruas_casas = []
        if self.var_esquina.get():
            for i in range(1, n_casas + 1):
                ruas_casas.append(self._get_rua_casa(i))
        rua2 = ruas_casas[1] if len(ruas_casas) >= 2 else ""

        dados_scpo = {
            "senha":       self.var_scpo_senha.get().strip(),
            "cep":         self.var_cep.get().strip(),
            "logradouro":  self.var_logradouro.get().strip(),
            "quadra":      quadra,
            "lote":        lote,
            "n_casas":     n_casas,
            "esquina":     self.var_esquina.get(),
            "rua2":        rua2,
            "ruas_casas":  ruas_casas,
            "data_inicio": self.var_scpo_data_inicio.get().strip(),
            "nome_obra":   _scpo_montar_nome_obra(
                               self.var_logradouro.get().strip(), quadra, lote),
            "observacao":  _scpo_montar_observacao(
                               self.var_logradouro.get().strip(), quadra, lote,
                               n_casas, self.var_esquina.get(),
                               rua2, ruas_casas),
        }

        # Resetar eventos
        self._scpo_evento_captcha = threading.Event()
        self._scpo_evento_envio   = threading.Event()

        # Desabilitar botão, limpar log, resetar progresso
        self.btn_scpo.config(state="disabled")
        self.btn_scpo_captcha.config(state="disabled")
        self.btn_scpo_envio.config(state="disabled")
        self.progress["value"] = 0
        self.var_status.set("SCPO: iniciando...")
        self.log("\n═══ INICIANDO SCPO ═══")

        def step_cb(pct, desc):
            self.after(0, lambda: (
                self.progress.config(value=pct),
                self.var_status.set(f"SCPO: {desc}")
            ))

        def done_cb(ok, msg):
            self.after(0, self._scpo_finalizar, ok, msg)

        threading.Thread(
            target=_scpo_executar,
            args=(dados_scpo, step_cb, self.log,
                  done_cb,
                  self._scpo_evento_captcha,
                  lambda: self.after(0, lambda: self.btn_scpo_captcha.config(state="normal")),
                  self._scpo_evento_envio,
                  lambda: self.after(0, lambda: self.btn_scpo_envio.config(state="normal"))),
            daemon=True
        ).start()

    def _scpo_finalizar(self, ok, msg):
        self.btn_scpo.config(state="normal")
        self.btn_scpo_captcha.config(state="disabled")
        self.btn_scpo_envio.config(state="disabled")
        if ok:
            messagebox.showinfo("SCPO", msg)
        else:
            messagebox.showerror("SCPO — Erro", msg)

    def _check_stop(self):
        """Levanta exceção se o usuário solicitou parada."""
        if self.stop_event.is_set():
            raise InterruptedError("Processamento interrompido pelo usuário")

    # ------------------------------------------------------------------
    # Thread de processamento
    # ------------------------------------------------------------------
    def _processar(self):
        try:
            eng_nome = self.var_engenheiro.get()
            eng_info = ENGENHEIROS[eng_nome]
            assinatura_path = asset(eng_info["assinatura"])

            if not os.path.exists(assinatura_path):
                raise FileNotFoundError(
                    f"Assinatura não encontrada: {assinatura_path}"
                )

            dados = {
                "engenheiro_nome": eng_nome,
                "cpf": self.var_cpf.get(),
                "crea": self.var_crea.get(),
                "art": self.var_art.get(),
                "contratante": self.var_contratante.get(),
                "logradouro": self._get_rua_casa(1),
                "quadra_lote": self.var_quadra_lote.get(),
                "bairro": self.var_bairro.get(),
                "cep": self.var_cep.get(),
                "cidade": self.var_cidade.get(),
                "uf": self.var_uf.get(),
                "assinatura_path": assinatura_path,
            }

            esgoto_sim = self.var_esgoto.get()

            _mg = {"Não se aplica": "nao_se_aplica", "Sim": "sim", "Não": "nao"}
            global GEMINADAS_CONDOMINIOS
            GEMINADAS_CONDOMINIOS = _mg.get(self.var_gem_cond.get(), "nao_se_aplica")
            qtd = self.var_qtd_casas.get()

            # Template Excel selecionado pelo usuário
            template_excel = self.var_memorial.get()
            if not template_excel or not os.path.exists(template_excel):
                raise FileNotFoundError("Selecione um arquivo Memorial válido.")

            # Pasta destino
            rua_qd_lt = f"{dados['logradouro']} {dados['quadra_lote']}".strip()
            rua_qd_lt = re.sub(r"[<>:\"/\\|?*]", "", rua_qd_lt)
            pasta_saida = Path.home() / "Downloads" / PASTA_DESTINO / rua_qd_lt
            pasta_saida.mkdir(parents=True, exist_ok=True)
            self.log(f"📁 Pasta destino: {pasta_saida}")

            total_etapas = qtd * 4
            etapa_atual = 0

            for i in range(1, qtd + 1):
                self._check_stop()
                self._set_status(f"Casa {i}/{qtd}...")
                self.log(f"\n═══ CASA {i} ═══")
                dados["logradouro"] = self._get_rua_casa(i)

                base_nome = f"CASA_{i:02d}"

                # 1. Word
                self._check_stop()
                self.log("• Gerando Declaração (Word)...")
                docx_path = pasta_saida / f"DECLARACAO_{base_nome}.docx"
                preencher_word(esgoto_sim, str(docx_path), dados, i, log=self.log)
                etapa_atual += 1
                self._set_progress(etapa_atual * 100 / total_etapas)

                # 2. Word → PDF
                self._check_stop()
                self.log("• Convertendo Declaração para PDF...")
                pdf_decl = pasta_saida / f"DECLARACAO_{base_nome}.pdf"
                _word_para_pdf(str(docx_path), str(pdf_decl), log=self.log)
                etapa_atual += 1
                self._set_progress(etapa_atual * 100 / total_etapas)

                # 3. Excel — preencher + checkboxes imagem (na mesma sessão COM)
                self._check_stop()
                self.log("• Preenchendo Memorial (Excel)...")
                xlsx_path = pasta_saida / f"MEMORIAL_{base_nome}.xlsx"
                _excel_preencher(
                    template_excel, str(xlsx_path), dados, i,
                    esgoto_sim, log=self.log,
                )
                etapa_atual += 1
                self._set_progress(etapa_atual * 100 / total_etapas)

                # 4. Excel → PDF
                self._check_stop()
                self.log("• Convertendo Memorial para PDF...")
                pdf_mem = pasta_saida / f"MEMORIAL_{base_nome}.pdf"
                _excel_para_pdf(str(xlsx_path), str(pdf_mem), log=self.log)
                etapa_atual += 1
                self._set_progress(etapa_atual * 100 / total_etapas)

            self._set_status("Concluído!")
            self._set_progress(100)
            self.log("\n✅ Todos os documentos gerados com sucesso!")
            self.log(f"📂 Pasta: {pasta_saida}")
            self.after(0, lambda: messagebox.showinfo(
                "Sucesso", f"Documentos gerados em:\n{pasta_saida}"))

        except InterruptedError:
            self.log("\n⛔ Processamento interrompido pelo usuário.")
            self._set_status("Interrompido.")
        except Exception as e:
            self.log(f"\n✗ ERRO: {e}")
            self.log(traceback.format_exc())
            self._set_status("Erro.")
            self.after(0, lambda: messagebox.showerror("Erro", str(e)))
        finally:
            self.processando = False
            self.stop_event.clear()
            self.after(0, lambda: self.btn_gerar.configure(state="normal"))
            self.after(0, lambda: self.btn_stop.configure(
                state="disabled", text="⛔ INTERROMPER"))


# ============================================================
# CALIBRADOR — janela Toplevel integrada ao app principal
# ============================================================

def _carregar_calibracao():
    """
    Lê ~/.bercan_config.json e aplica os valores calibrados nas
    constantes globais CHK* e ASSINATURA_EXCEL_*.
    Se o arquivo não existir ou a chave estiver ausente, mantém o default.
    """
    cfg = _config_carregar()
    if not cfg.get("calibrado"):
        return  # ainda não foi calibrado — usa defaults

    global CHK1_ANCORA, CHK1_OFF_X, CHK1_OFF_Y, CHK1_LARGURA, CHK1_ALTURA
    global CHK2_ANCORA, CHK2_OFF_X, CHK2_OFF_Y, CHK2_LARGURA, CHK2_ALTURA
    global CHK3_ANCORA, CHK3_OFF_X, CHK3_OFF_Y, CHK3_LARGURA, CHK3_ALTURA
    global CHK4_ANCORA, CHK4_OFF_X, CHK4_OFF_Y, CHK4_LARGURA, CHK4_ALTURA
    global CHK_LOT_NSA_ANCORA, CHK_LOT_NSA_OFF_X, CHK_LOT_NSA_OFF_Y
    global CHK_LOT_NSA_LARGURA, CHK_LOT_NSA_ALTURA
    global ASSINATURA_EXCEL_ANCORA, ASSINATURA_EXCEL_OFFSET_X_PT
    global ASSINATURA_EXCEL_OFFSET_Y_PT, ASSINATURA_EXCEL_LARGURA_PT
    global ASSINATURA_EXCEL_ALTURA_PT

    def _i(k, fallback): return int(cfg.get(k, fallback))
    def _s(k, fallback): return str(cfg.get(k, fallback))

    CHK1_ANCORA  = _s("chk1_ancora",  CHK1_ANCORA)
    CHK1_OFF_X   = _i("chk1_off_x",   CHK1_OFF_X)
    CHK1_OFF_Y   = _i("chk1_off_y",   CHK1_OFF_Y)
    CHK1_LARGURA = _i("chk1_larg",    CHK1_LARGURA)
    CHK1_ALTURA  = _i("chk1_alt",     CHK1_ALTURA)

    CHK2_ANCORA  = _s("chk2_ancora",  CHK2_ANCORA)
    CHK2_OFF_X   = _i("chk2_off_x",   CHK2_OFF_X)
    CHK2_OFF_Y   = _i("chk2_off_y",   CHK2_OFF_Y)
    CHK2_LARGURA = _i("chk2_larg",    CHK2_LARGURA)
    CHK2_ALTURA  = _i("chk2_alt",     CHK2_ALTURA)

    CHK3_ANCORA  = _s("chk3_ancora",  CHK3_ANCORA)
    CHK3_OFF_X   = _i("chk3_off_x",   CHK3_OFF_X)
    CHK3_OFF_Y   = _i("chk3_off_y",   CHK3_OFF_Y)
    CHK3_LARGURA = _i("chk3_larg",    CHK3_LARGURA)
    CHK3_ALTURA  = _i("chk3_alt",     CHK3_ALTURA)

    CHK4_ANCORA  = _s("chk4_ancora",  CHK4_ANCORA)
    CHK4_OFF_X   = _i("chk4_off_x",   CHK4_OFF_X)
    CHK4_OFF_Y   = _i("chk4_off_y",   CHK4_OFF_Y)
    CHK4_LARGURA = _i("chk4_larg",    CHK4_LARGURA)
    CHK4_ALTURA  = _i("chk4_alt",     CHK4_ALTURA)

    CHK_LOT_NSA_ANCORA  = _s("chk_lot_ancora", CHK_LOT_NSA_ANCORA)
    CHK_LOT_NSA_OFF_X   = _i("chk_lot_off_x",  CHK_LOT_NSA_OFF_X)
    CHK_LOT_NSA_OFF_Y   = _i("chk_lot_off_y",  CHK_LOT_NSA_OFF_Y)
    CHK_LOT_NSA_LARGURA = _i("chk_lot_larg",   CHK_LOT_NSA_LARGURA)
    CHK_LOT_NSA_ALTURA  = _i("chk_lot_alt",    CHK_LOT_NSA_ALTURA)

    ASSINATURA_EXCEL_ANCORA      = _s("ass_ancora",  ASSINATURA_EXCEL_ANCORA)
    ASSINATURA_EXCEL_OFFSET_X_PT = _i("ass_off_x",   ASSINATURA_EXCEL_OFFSET_X_PT)
    ASSINATURA_EXCEL_OFFSET_Y_PT = _i("ass_off_y",   ASSINATURA_EXCEL_OFFSET_Y_PT)
    ASSINATURA_EXCEL_LARGURA_PT  = _i("ass_larg",    ASSINATURA_EXCEL_LARGURA_PT)
    ASSINATURA_EXCEL_ALTURA_PT   = _i("ass_alt",     ASSINATURA_EXCEL_ALTURA_PT)


class JanelaCalibrador(tk.Toplevel):
    """
    Calibrador do Memorial — abre como janela filha do app principal.
    Permite ajustar posição/tamanho dos checkboxes e assinatura,
    gera um PDF de preview e salva os valores no config JSON.
    """

    # Mapa de estados para calibrar
    ESTADOS = {
        1: ("Esgoto — SIM",               "#2e86de"),
        2: ("Esgoto — NÃO",               "#e74c3c"),
        3: ("Condomínio — SIM",           "#27ae60"),
        4: ("Condomínio — Não se aplica", "#e67e22"),
    }

    def __init__(self, parent, memorial_path=None):
        super().__init__(parent)
        self.title("⚙ Calibrador do Memorial — Morais Engenharia")
        self.geometry("860x780")
        self.configure(bg=COR_FUNDO)
        self.resizable(True, True)
        self.grab_set()  # modal

        self._vars = {}
        self._estado_atual = 1
        self._memorial_path = memorial_path

        self._criar_ui()

        # Pré-carregar valores do config se existir
        self._carregar_do_config()

        # Pré-preencher memorial se já estava selecionado no app
        if memorial_path:
            self._vars["memorial"].set(memorial_path)

    # ── Construção da UI ─────────────────────────────────────────────

    def _criar_ui(self):
        # Cabeçalho
        hdr = tk.Frame(self, bg=COR_FUNDO, pady=10)
        hdr.pack(fill="x", padx=20)
        tk.Label(hdr, text="CALIBRADOR DO MEMORIAL",
                 font=("Segoe UI", 14, "bold"),
                 fg=COR_TEXTO, bg=COR_FUNDO).pack(anchor="w")
        tk.Label(hdr,
                 text="Ajuste checkboxes e assinatura • preview em PDF • salva automaticamente",
                 font=("Segoe UI", 9), fg=COR_TEXTO_SEC, bg=COR_FUNDO).pack(anchor="w")
        tk.Frame(self, bg="#2a3f55", height=1).pack(fill="x")

        # Rodapé com botões
        rodape = tk.Frame(self, bg=COR_FUNDO, pady=8)
        rodape.pack(side="bottom", fill="x", padx=20)
        tk.Frame(rodape, bg="#2a3f55", height=1).pack(fill="x", pady=(0, 8))

        self.btn_preview = tk.Button(
            rodape, text="⚡ GERAR PREVIEW (abre PDF)",
            command=self._iniciar_preview,
            bg=COR_BOTAO, fg=COR_TEXTO, relief="flat",
            font=("Segoe UI", 11, "bold"), pady=8,
        )
        self.btn_preview.pack(fill="x", pady=(0, 4))

        tk.Button(
            rodape, text="💾 SALVAR CALIBRAÇÃO",
            command=self._salvar,
            bg="#27ae60", fg=COR_TEXTO, relief="flat",
            font=("Segoe UI", 11, "bold"), pady=8,
        ).pack(fill="x")

        self.var_status = tk.StringVar(value="Aguardando...")
        tk.Label(rodape, textvariable=self.var_status,
                 bg=COR_FUNDO, fg=COR_TEXTO_SEC,
                 font=("Segoe UI", 9)).pack(anchor="w", pady=(4, 0))

        # Corpo em duas colunas
        body = tk.Frame(self, bg=COR_FUNDO)
        body.pack(fill="both", expand=True, padx=20, pady=8)

        # Coluna esquerda com scroll — evita corte dos campos inferiores
        outer_esq = tk.Frame(body, bg=COR_FUNDO, width=390)
        outer_esq.pack(side="left", fill="y", padx=(0, 12))
        outer_esq.pack_propagate(False)
        canvas_esq = tk.Canvas(outer_esq, bg=COR_FUNDO, highlightthickness=0)
        sb_esq = tk.Scrollbar(outer_esq, orient="vertical", command=canvas_esq.yview)
        canvas_esq.configure(yscrollcommand=sb_esq.set)
        sb_esq.pack(side="right", fill="y")
        canvas_esq.pack(side="left", fill="both", expand=True)
        col_esq = tk.Frame(canvas_esq, bg=COR_FUNDO)
        wid_esq = canvas_esq.create_window((0, 0), window=col_esq, anchor="nw")
        col_esq.bind("<Configure>", lambda e: (
            canvas_esq.configure(scrollregion=canvas_esq.bbox("all")),
            canvas_esq.itemconfig(wid_esq, width=canvas_esq.winfo_width()),
        ))
        canvas_esq.bind("<MouseWheel>",
            lambda e: canvas_esq.yview_scroll(int(-1*(e.delta/120)), "units"))
        col_esq.bind("<MouseWheel>",
            lambda e: canvas_esq.yview_scroll(int(-1*(e.delta/120)), "units"))

        col_dir = tk.Frame(body, bg=COR_FUNDO)
        col_dir.pack(side="left", fill="both", expand=True)

        # ── Coluna esquerda: controles ───────────────────────────────
        self._label_sec(col_esq, "ARQUIVO MEMORIAL")
        self._vars["memorial"] = tk.StringVar()
        self._campo_arquivo(col_esq, self._vars["memorial"],
                            "Selecionar .xls/.xlsx")

        self._label_sec(col_esq, "ASSINATURA (opcional)")
        self._vars["assinatura"] = tk.StringVar()
        self._campo_arquivo(col_esq, self._vars["assinatura"],
                            "Selecionar imagem de assinatura")

        self._label_sec(col_esq, "ASSINATURA — POSIÇÃO")
        self._spinbox(col_esq, "Âncora",   "ass_ancora",  "AE72", texto=True)
        self._spinbox(col_esq, "Offset X", "ass_off_x",   10, -200, 200)
        self._spinbox(col_esq, "Offset Y", "ass_off_y",   -5, -200, 200)
        self._spinbox(col_esq, "Largura",  "ass_larg",    170,  10, 500)
        self._spinbox(col_esq, "Altura",   "ass_alt",     55,   5, 300)

        self._label_sec(col_esq, "CHECKBOX — ESTADO")
        fr_est = tk.Frame(col_esq, bg=COR_FUNDO)
        fr_est.pack(fill="x", pady=(0, 4))
        self.var_estado = tk.IntVar(value=1)
        for n, (lbl, cor) in self.ESTADOS.items():
            tk.Radiobutton(
                fr_est, text=lbl, variable=self.var_estado, value=n,
                command=self._trocar_estado,
                bg=COR_FUNDO, fg=cor, selectcolor=COR_CAMPO,
                activebackground=COR_FUNDO, activeforeground=cor,
                font=("Segoe UI", 9),
            ).pack(anchor="w")

        self.lbl_estado = tk.Label(col_esq,
                                   text="Estado 1 — Esgoto — SIM",
                                   fg="#2e86de", bg=COR_FUNDO,
                                   font=("Segoe UI", 9, "bold"))
        self.lbl_estado.pack(anchor="w", pady=(2, 4))

        self._label_sec(col_esq, "CHECKBOX — POSIÇÃO")
        self._spinbox(col_esq, "Âncora",   "chk_ancora", "AM70", texto=True)
        self._spinbox(col_esq, "Offset X", "chk_off_x",   10, -200, 200)
        self._spinbox(col_esq, "Offset Y", "chk_off_y",    3, -200, 200)
        self._spinbox(col_esq, "Largura",  "chk_larg",     4,   1, 100)
        self._spinbox(col_esq, "Altura",   "chk_alt",      5,   1, 100)

        # ── Coluna direita: log ──────────────────────────────────────
        self._label_sec(col_dir, "LOG")
        self.txt_log = tk.Text(
            col_dir, bg=COR_LOG_FUNDO, fg="#7ec8a0",
            font=("Consolas", 9), relief="flat",
            state="disabled",
        )
        sb = tk.Scrollbar(col_dir, command=self.txt_log.yview)
        self.txt_log.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        self.txt_log.pack(fill="both", expand=True)

    # ── Widgets helper ────────────────────────────────────────────────

    def _label_sec(self, p, txt):
        tk.Label(p, text=txt, font=("Segoe UI", 9, "bold"),
                 fg=COR_TEXTO, bg=COR_FUNDO).pack(anchor="w", pady=(10, 2))

    def _campo_arquivo(self, p, var, hint):
        fr = tk.Frame(p, bg=COR_FUNDO)
        fr.pack(fill="x", pady=(0, 4))
        tk.Entry(fr, textvariable=var, bg=COR_CAMPO, fg=COR_TEXTO,
                 insertbackground=COR_TEXTO, relief="flat",
                 font=("Segoe UI", 9)).pack(side="left", fill="x", expand=True)
        tk.Button(fr, text="📁",
                  command=lambda: var.set(
                      filedialog.askopenfilename(
                          filetypes=[("Excel/Imagem",
                                      "*.xls *.xlsx *.png *.jpg *.jpeg")]) or var.get()),
                  bg=COR_BOTAO, fg=COR_TEXTO, relief="flat",
                  font=("Segoe UI", 9)).pack(side="right", padx=(4, 0))

    def _spinbox(self, p, label, key, default, mn=-999, mx=9999, texto=False):
        fr = tk.Frame(p, bg=COR_FUNDO)
        fr.pack(fill="x", pady=1)
        tk.Label(fr, text=label, width=10, anchor="w",
                 fg=COR_TEXTO_SEC, bg=COR_FUNDO,
                 font=("Segoe UI", 9)).pack(side="left")
        var = tk.StringVar(value=str(default))
        self._vars[key] = var
        if texto:
            tk.Entry(fr, textvariable=var, width=10, bg=COR_CAMPO, fg=COR_TEXTO,
                     insertbackground=COR_TEXTO, relief="flat",
                     font=("Segoe UI", 10, "bold")).pack(side="left")
        else:
            tk.Spinbox(fr, textvariable=var, from_=mn, to=mx, width=7,
                       bg=COR_CAMPO, fg=COR_TEXTO, buttonbackground=COR_CAMPO,
                       relief="flat", insertbackground=COR_TEXTO,
                       font=("Segoe UI", 10, "bold")).pack(side="left")
            tk.Button(fr, text="-5",
                      command=lambda v=var: self._nudge(v, -5),
                      bg="#2a3f55", fg=COR_TEXTO_SEC, relief="flat",
                      font=("Segoe UI", 8), padx=3).pack(side="left", padx=(6, 1))
            tk.Button(fr, text="+5",
                      command=lambda v=var: self._nudge(v, +5),
                      bg="#2a3f55", fg=COR_TEXTO_SEC, relief="flat",
                      font=("Segoe UI", 8), padx=3).pack(side="left", padx=1)

    def _nudge(self, var, d):
        try: var.set(str(int(var.get()) + d))
        except: pass

    # ── Estado dos checkboxes ─────────────────────────────────────────

    # Armazena valores por estado (1-4)
    _estado_vals = {}

    def _trocar_estado(self):
        self._salvar_estado_atual()
        n = self.var_estado.get()
        self._estado_atual = n
        lbl, cor = self.ESTADOS[n]
        self.lbl_estado.configure(text=f"Estado {n} — {lbl}", fg=cor)
        self._carregar_estado(n)

    def _salvar_estado_atual(self):
        n = self._estado_atual
        v = self._vars
        self._estado_vals[n] = {
            "ancora": v["chk_ancora"].get().strip(),
            "off_x":  self._int("chk_off_x", 0),
            "off_y":  self._int("chk_off_y", 0),
            "larg":   self._int("chk_larg",   4),
            "alt":    self._int("chk_alt",    5),
        }

    def _carregar_estado(self, n):
        if n not in self._estado_vals:
            return
        d = self._estado_vals[n]
        self._vars["chk_ancora"].set(d["ancora"])
        self._vars["chk_off_x"].set(str(d["off_x"]))
        self._vars["chk_off_y"].set(str(d["off_y"]))
        self._vars["chk_larg"].set(str(d["larg"]))
        self._vars["chk_alt"].set(str(d["alt"]))

    def _int(self, key, fallback=0):
        try: return int(self._vars[key].get())
        except: return fallback

    # ── Config ────────────────────────────────────────────────────────

    def _carregar_do_config(self):
        """Preenche os campos com valores já salvos no JSON."""
        cfg = _config_carregar()
        if not cfg.get("calibrado"):
            # Sem calibração salva — inicializa estado_vals com defaults
            self._estado_vals = {
                1: {"ancora": CHK1_ANCORA,  "off_x": CHK1_OFF_X,  "off_y": CHK1_OFF_Y,  "larg": CHK1_LARGURA,  "alt": CHK1_ALTURA},
                2: {"ancora": CHK2_ANCORA,  "off_x": CHK2_OFF_X,  "off_y": CHK2_OFF_Y,  "larg": CHK2_LARGURA,  "alt": CHK2_ALTURA},
                3: {"ancora": CHK3_ANCORA,  "off_x": CHK3_OFF_X,  "off_y": CHK3_OFF_Y,  "larg": CHK3_LARGURA,  "alt": CHK3_ALTURA},
                4: {"ancora": CHK4_ANCORA,  "off_x": CHK4_OFF_X,  "off_y": CHK4_OFF_Y,  "larg": CHK4_LARGURA,  "alt": CHK4_ALTURA},
            }
        else:
            self._estado_vals = {
                1: {"ancora": cfg.get("chk1_ancora", CHK1_ANCORA), "off_x": cfg.get("chk1_off_x", CHK1_OFF_X), "off_y": cfg.get("chk1_off_y", CHK1_OFF_Y), "larg": cfg.get("chk1_larg", CHK1_LARGURA), "alt": cfg.get("chk1_alt", CHK1_ALTURA)},
                2: {"ancora": cfg.get("chk2_ancora", CHK2_ANCORA), "off_x": cfg.get("chk2_off_x", CHK2_OFF_X), "off_y": cfg.get("chk2_off_y", CHK2_OFF_Y), "larg": cfg.get("chk2_larg", CHK2_LARGURA), "alt": cfg.get("chk2_alt", CHK2_ALTURA)},
                3: {"ancora": cfg.get("chk3_ancora", CHK3_ANCORA), "off_x": cfg.get("chk3_off_x", CHK3_OFF_X), "off_y": cfg.get("chk3_off_y", CHK3_OFF_Y), "larg": cfg.get("chk3_larg", CHK3_LARGURA), "alt": cfg.get("chk3_alt", CHK3_ALTURA)},
                4: {"ancora": cfg.get("chk4_ancora", CHK4_ANCORA), "off_x": cfg.get("chk4_off_x", CHK4_OFF_X), "off_y": cfg.get("chk4_off_y", CHK4_OFF_Y), "larg": cfg.get("chk4_larg", CHK4_LARGURA), "alt": cfg.get("chk4_alt", CHK4_ALTURA)},
            }
            self._vars["ass_ancora"].set(cfg.get("ass_ancora", "AE72"))
            self._vars["ass_off_x"].set(str(cfg.get("ass_off_x", 10)))
            self._vars["ass_off_y"].set(str(cfg.get("ass_off_y", -5)))
            self._vars["ass_larg"].set(str(cfg.get("ass_larg", 170)))
            self._vars["ass_alt"].set(str(cfg.get("ass_alt", 55)))

        # Carregar estado 1 nos campos visíveis
        self._carregar_estado(1)

    def _salvar(self):
        """Salva todos os valores calibrados no config JSON e atualiza as globais."""
        self._salvar_estado_atual()
        payload = {
            "calibrado": True,
            "ass_ancora": self._vars["ass_ancora"].get().strip(),
            "ass_off_x":  self._int("ass_off_x"),
            "ass_off_y":  self._int("ass_off_y"),
            "ass_larg":   self._int("ass_larg"),
            "ass_alt":    self._int("ass_alt"),
        }
        keys = ["ancora", "off_x", "off_y", "larg", "alt"]
        for n in range(1, 5):
            d = self._estado_vals.get(n, {})
            payload[f"chk{n}_ancora"] = d.get("ancora", "")
            payload[f"chk{n}_off_x"]  = d.get("off_x", 0)
            payload[f"chk{n}_off_y"]  = d.get("off_y", 0)
            payload[f"chk{n}_larg"]   = d.get("larg", 4)
            payload[f"chk{n}_alt"]    = d.get("alt", 5)

        _config_salvar(payload)
        _carregar_calibracao()  # aplica imediatamente nas globais
        self.var_status.set("✅ Calibração salva! Próxima geração já usa os novos valores.")
        self._log("✅ Valores salvos em ~/.bercan_config.json e aplicados.")

    # ── Preview ───────────────────────────────────────────────────────

    def _iniciar_preview(self):
        memorial = self._vars["memorial"].get().strip()
        if not memorial or not os.path.exists(memorial):
            messagebox.showerror("Erro", "Selecione um Memorial válido.", parent=self)
            return
        if "PREVIEW_CALIBRADOR" in os.path.basename(memorial):
            messagebox.showerror("Arquivo inválido",
                                 "Selecione o memorial ORIGINAL, não o preview.",
                                 parent=self)
            return

        self._salvar_estado_atual()
        self.btn_preview.configure(state="disabled", text="⏳ Gerando preview...")
        self.var_status.set("Processando...")

        estado   = self.var_estado.get()
        ass_img  = self._vars["assinatura"].get().strip() or None
        saida    = str(Path(memorial).parent / "PREVIEW_CALIBRADOR.xlsx")

        threading.Thread(
            target=self._worker_preview,
            args=(memorial, ass_img, estado, saida),
            daemon=True,
        ).start()

    def _worker_preview(self, memorial, ass_img, estado, saida):
        import pythoncom, win32com.client, time as _time

        def log(msg):
            self.after(0, self._log, msg)

        pythoncom.CoInitialize()
        xl = None; wb = None
        try:
            # Copiar template virgem para preview
            log("• Copiando memorial...")
            shutil.copy2(memorial, saida)

            log("• Abrindo no Excel...")
            xl = win32com.client.Dispatch("Excel.Application")
            try: xl.Visible = False
            except: pass
            try: xl.DisplayAlerts = False
            except: pass
            wb = xl.Workbooks.Open(os.path.abspath(saida))
            try: ws = wb.Worksheets("ElemConstrutivos")
            except: ws = wb.Worksheets(1)

            # Assinatura
            if ass_img and os.path.exists(ass_img):
                log("• Inserindo assinatura...")
                cell = ws.Range(self._vars["ass_ancora"].get().strip())
                ws.Shapes.AddPicture(
                    os.path.abspath(ass_img), False, True,
                    cell.Left + self._int("ass_off_x"),
                    cell.Top  + self._int("ass_off_y"),
                    self._int("ass_larg"),
                    self._int("ass_alt"),
                )

            # Checkbox do estado atual
            d = self._estado_vals.get(estado, {})
            log(f"• Inserindo checkbox estado {estado} ({self.ESTADOS[estado][0]})...")
            q = _quadrado_preto_temp()
            try:
                cell = ws.Range(d.get("ancora", "AM70"))
                ws.Shapes.AddPicture(
                    os.path.abspath(q), False, True,
                    cell.Left + d.get("off_x", 0),
                    cell.Top  + d.get("off_y", 0),
                    d.get("larg", 4),
                    d.get("alt", 5),
                )
            finally:
                try: os.unlink(q)
                except: pass

            # Exportar PDF
            wb.Save()
            pdf = saida.replace(".xlsx", ".pdf")
            ws.PageSetup.Zoom = False
            ws.PageSetup.FitToPagesWide = 1
            ws.PageSetup.FitToPagesTall = 1
            ws.ExportAsFixedFormat(0, os.path.abspath(pdf), 0, True, False)
            log(f"✅ PDF gerado: {os.path.basename(pdf)}")
            log("─" * 40)
            log(f"Estado calibrado: {self.ESTADOS[estado][0]}")
            log(f"  ancora={d.get('ancora')} off=({d.get('off_x')},{d.get('off_y')}) "
                f"tam={d.get('larg')}x{d.get('alt')}")

            try: os.startfile(pdf)
            except: pass

            self.after(0, self.var_status.set, "✅ Preview gerado — verifique o PDF!")

        except Exception as e:
            log(f"✗ ERRO: {e}\n{traceback.format_exc()}")
            self.after(0, self.var_status.set, f"✗ {e}")
        finally:
            if wb:
                try: wb.Close(SaveChanges=False)
                except: pass
            if xl:
                try: xl.Quit()
                except: pass
            pythoncom.CoUninitialize()
            self.after(0, self.btn_preview.configure,
                       {"state": "normal",
                        "text": "⚡ GERAR PREVIEW (abre PDF)"})

    def _log(self, msg):
        self.txt_log.configure(state="normal")
        self.txt_log.insert("end", msg + "\n")
        self.txt_log.see("end")
        self.txt_log.configure(state="disabled")


# ============================================================
# ENTRY POINT
# ============================================================
if __name__ == "__main__":
    _carregar_calibracao()  # aplica config salvo antes de abrir a janela
    App().mainloop()
